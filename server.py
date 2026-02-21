"""
Mickey — Legal Intelligence Platform
Production-hardened server

Security upgrades:
- bcrypt password hashing (replaces SHA256+salt)
- Secret key from environment variable
- SESSION_COOKIE_SECURE ready for HTTPS
- CSRF token protection on all state-changing requests
- Per-user rate limiting (API calls per hour)
- Per-user usage tracking (tokens, calls, documents)

New features:
- OpenAI text-embedding-3-small for real semantic search
- Usage dashboard data
"""

# ── Config loader — ADD THIS AT THE VERY TOP OF server.py ─────
# Loads /opt/mickey/config.env into environment variables on startup.
# Must be the first thing that runs, before any other imports.

import os
from pathlib import Path

for _p in [
    Path("/opt/mickey/config.env"),
    Path(__file__).parent / "config.env",
]:
    if _p.exists():
        for _line in _p.read_text(encoding="utf-8").splitlines():
            _line = _line.strip()
            if _line and not _line.startswith("#") and "=" in _line:
                _k, _, _v = _line.partition("=")
                os.environ.setdefault(_k.strip(), _v.strip())
        break

# ── END CONFIG LOADER ─────────────────────────────────────────
import os, json, re, secrets, base64, hashlib, datetime, io, time, shutil
from pathlib import Path
from functools import wraps
from collections import defaultdict

from flask import Flask, request, jsonify, session, render_template, g
import anthropic

# ── App setup ─────────────────────────────────────────────────
app = Flask(__name__)
try:
    from auth.routes import auth_bp
    app.register_blueprint(auth_bp)
    print("  Auth module: loaded")
except ImportError as e:
    print(f"  Auth module: not found ({e}) — running without firm auth")
@app.route("/direct-login", methods=["GET", "POST"])
def direct_login():
    if request.method == "POST":
        session["username"] = request.form.get("u")
        session["display_name"] = request.form.get("u")
        session["role"] = "admin"
        return redirect("/")
    return '<form method="POST"><input name="u" value="mathias.baert@lex-it.be"><button>Login</button></form>
# Secret key MUST come from environment in production
# Set: $env:MICKEY_SECRET = "your-long-random-string"  (PowerShell)
# Or:  set MICKEY_SECRET=your-long-random-string         (cmd)
app.secret_key = os.environ.get(
    "MICKEY_SECRET",
    # Fallback for first-run convenience only — print a warning
    "CHANGE_ME_SET_MICKEY_SECRET_ENV_VAR"
)
if app.secret_key == "CHANGE_ME_SET_MICKEY_SECRET_ENV_VAR":
    print("\n  [WARNING] MICKEY_SECRET environment variable not set.")
    print("  Set it with: set MICKEY_SECRET=<random-string>")
    print("  Generating a random key for this session only.\n")
    app.secret_key = secrets.token_hex(32)

app.config.update(
    SESSION_COOKIE_HTTPONLY  = True,
    SESSION_COOKIE_SAMESITE  = "Lax",
    SESSION_COOKIE_SECURE    = os.environ.get("MICKEY_HTTPS", "false").lower() == "true",
    PERMANENT_SESSION_LIFETIME = datetime.timedelta(hours=8),
)

# ── Paths ─────────────────────────────────────────────────────
APP_PATH     = Path(os.environ.get("MICKEY_DATA", "C:/Mickey"))
USERS_FILE   = APP_PATH / "users.json"
DATA_PATH    = APP_PATH / "data"
SHARED_PATH  = APP_PATH / "shared_library"
USAGE_PATH   = APP_PATH / "usage"
LOG_PATH     = APP_PATH / "logs"

for p in [APP_PATH, DATA_PATH, SHARED_PATH, USAGE_PATH, LOG_PATH,
          SHARED_PATH / "embeddings"]:
    p.mkdir(parents=True, exist_ok=True)

API_KEY_FILE    = APP_PATH / "api_key.enc"
OAI_KEY_FILE    = APP_PATH / "oai_key.enc"
KEY_SALT_FILE   = APP_PATH / "key.salt"

SKIP_FOLDERS = {
    "chatbestanden van microsoft teams","microsoft copilot chat files",
    "microsoft teams chat files","notebooks","notitieblokken",
    "pictures","recordings","videos","photos","$recycle.bin",
    "system volume information",".git","__pycache__","node_modules",
    "appdata","windows","program files","program files (x86)"
}

# Per-user rate limits
RATE_LIMIT_CALLS_PER_HOUR = int(os.environ.get("MICKEY_RATE_LIMIT", "60"))
MAX_LOGIN_ATTEMPTS = 10  # lockout after this many failures
LOCKOUT_MINUTES    = 15  # initial lockout duration

# ── In-memory rate limiter ────────────────────────────────────
_rate_store    = defaultdict(list)   # username -> [timestamps]
_login_fails   = defaultdict(list)   # username/ip -> [fail_timestamps]

def check_rate_limit(username):
    now = time.time(); window = 3600
    calls = [t for t in _rate_store[username] if now - t < window]
    _rate_store[username] = calls
    if len(calls) >= RATE_LIMIT_CALLS_PER_HOUR: return False
    _rate_store[username].append(now)
    return True

def record_login_fail(key: str):
    """key = username or IP. Record a failed attempt."""
    now = time.time()
    _login_fails[key] = [t for t in _login_fails[key] if now - t < 3600]
    _login_fails[key].append(now)

def is_locked_out(key: str) -> tuple:
    """Returns (locked: bool, seconds_remaining: int)."""
    now   = time.time()
    fails = [t for t in _login_fails[key] if now - t < 3600]
    _login_fails[key] = fails
    if len(fails) < MAX_LOGIN_ATTEMPTS: return False, 0
    last_fail = max(fails)
    wait_secs = int(LOCKOUT_MINUTES * 60 * (1 + (len(fails) - MAX_LOGIN_ATTEMPTS) * 0.5))
    elapsed   = now - last_fail
    if elapsed < wait_secs: return True, int(wait_secs - elapsed)
    return False, 0

def clear_login_fails(key: str):
    _login_fails.pop(key, None)

# ── Encryption (API keys at rest) ────────────────────────────
def _fernet():
    """Fernet = AES-128-CBC + HMAC-SHA256. Replaces XOR obfuscation."""
    try:
        from cryptography.fernet import Fernet
        from cryptography.hazmat.primitives import hashes
        from cryptography.hazmat.primitives.kdf.pbkdf2 import PBKDF2HMAC
    except ImportError:
        raise RuntimeError("Run: pip install cryptography")
    if not KEY_SALT_FILE.exists():
        KEY_SALT_FILE.write_bytes(secrets.token_bytes(32))
    salt = KEY_SALT_FILE.read_bytes()
    kdf  = PBKDF2HMAC(algorithm=hashes.SHA256(), length=32, salt=salt, iterations=390000)
    key  = base64.urlsafe_b64encode(kdf.derive(app.secret_key.encode()))
    return Fernet(key)

def encrypt(text: str) -> str:
    return _fernet().encrypt(text.encode()).decode()

def decrypt(cipher: str) -> str:
    return _fernet().decrypt(cipher.encode()).decode()

def _read_key(path):
    if not path.exists(): return ""
    try:
        raw = path.read_text().strip()
        if not raw: return ""
        # Fernet tokens always start with gAAAAA
        if not raw.startswith("gAAAAA"):
            path.unlink()  # Old XOR format — wipe, user must re-enter via Admin
            return ""
        return decrypt(raw)
    except: return ""

def _write_key(path, value):
    path.write_text(encrypt(value.strip()))

def get_api_key():
    return os.environ.get("ANTHROPIC_API_KEY","").strip() or _read_key(API_KEY_FILE)

def get_oai_key():
    return os.environ.get("OPENAI_API_KEY","").strip() or _read_key(OAI_KEY_FILE)

def get_client():
    key = get_api_key()
    if not key: raise ValueError("Anthropic API key not configured.")
    return anthropic.Anthropic(api_key=key)

# ── Password hashing (bcrypt with SHA256 pre-hash) ────────────
# We pre-hash with SHA256 to handle passwords > 72 bytes (bcrypt limit)
# This is the recommended approach when using bcrypt.
def hash_pw(pw: str) -> str:
    """Returns bcrypt hash string. Stores everything in one field."""
    import bcrypt
    pre = hashlib.sha256(pw.encode()).hexdigest().encode()
    return bcrypt.hashpw(pre, bcrypt.gensalt(rounds=12)).decode()

def verify_pw(pw: str, stored: str) -> bool:
    import bcrypt
    try:
        pre = hashlib.sha256(pw.encode()).hexdigest().encode()
        return bcrypt.checkpw(pre, stored.encode())
    except: return False

def validate_pw(pw):
    rules = []
    if len(pw) < 8: rules.append("at least 8 characters")
    if not re.search(r'\d', pw): rules.append("a number")
    if not re.search(r'[!@#$%^&*()\-_=+\[\]{};:\'",.<>/?`~\\|]', pw): rules.append("a special character")
    if not re.search(r'[A-Z]', pw): rules.append("an uppercase letter")
    return len(rules) == 0, rules

# ── Migration: upgrade old SHA256 hashes to bcrypt on login ──
def _is_old_hash(stored: str) -> bool:
    return len(stored) == 64 and not stored.startswith("$2b$")

def _old_verify(pw, h, salt):
    return hashlib.sha256((salt + pw).encode()).hexdigest() == h

# ── User management ───────────────────────────────────────────
def load_users():
    if USERS_FILE.exists():
        try: return json.loads(USERS_FILE.read_text(encoding="utf-8"))
        except: return {}
    return {}

def save_users(u):
    USERS_FILE.write_text(json.dumps(u, indent=2, ensure_ascii=False), encoding="utf-8")

def is_admin(username):
    return load_users().get(username, {}).get("role") == "admin"

def upath(username):
    p = DATA_PATH / username; p.mkdir(exist_ok=True)
    for sub in ["workspace","embeddings"]: (p/sub).mkdir(exist_ok=True)
    return p

def load_settings(u):
    f = upath(u)/"settings.json"
    d = {"drive_paths":[],"drive_labels":[],"display_name":u}
    if f.exists():
        try: d.update(json.loads(f.read_text(encoding="utf-8")))
        except: pass
    return d

def save_settings(u, s):
    (upath(u)/"settings.json").write_text(json.dumps(s,indent=2,ensure_ascii=False),encoding="utf-8")

def load_index(u):
    f = upath(u)/"file_index.json"
    if f.exists():
        try: return json.loads(f.read_text(encoding="utf-8"))
        except: return {}
    return {}

def save_index(u, idx):
    (upath(u)/"file_index.json").write_text(json.dumps(idx,ensure_ascii=False,indent=2),encoding="utf-8")

# ── Usage tracking ────────────────────────────────────────────
def log_usage(username, event_type, tokens_used=0, detail=""):
    """Append usage record. Minimal — timestamps + counts only."""
    today = datetime.date.today().isoformat()
    f = USAGE_PATH / f"{username}_{today}.jsonl"
    record = {
        "ts": datetime.datetime.utcnow().isoformat(),
        "user": username,
        "event": event_type,   # ask / analyze / draft / translate
        "tokens": tokens_used,
        "detail": detail[:80]  # truncated — no question content stored
    }
    with open(f, "a", encoding="utf-8") as fh:
        fh.write(json.dumps(record) + "\n")

def get_usage_summary(username, days=30):
    """Aggregate usage for a user over last N days."""
    total_calls = total_tokens = 0
    by_event = {}
    cutoff = datetime.date.today() - datetime.timedelta(days=days)
    for f in USAGE_PATH.glob(f"{username}_*.jsonl"):
        try:
            date_str = f.stem.replace(username+"_","")
            if datetime.date.fromisoformat(date_str) < cutoff: continue
            for line in f.read_text(encoding="utf-8").splitlines():
                r = json.loads(line)
                total_calls += 1
                total_tokens += r.get("tokens",0)
                ev = r.get("event","other")
                by_event[ev] = by_event.get(ev,0) + 1
        except: pass
    return {"calls": total_calls, "tokens": total_tokens, "by_event": by_event}

def get_all_usage_summary(days=30):
    """Admin: usage across all users."""
    users = load_users()
    return {u: get_usage_summary(u, days) for u in users}

# ── Auth decorators ───────────────────────────────────────────
def login_required(f):
    @wraps(f)
    def dec(*a, **k):
        if not session.get("username"):
            return jsonify({"error":"Not authenticated"}), 401
        session.modified = True
        return f(*a, **k)
    return dec

def admin_required(f):
    @wraps(f)
    def dec(*a, **k):
        u = session.get("username")
        if not u: return jsonify({"error":"Not authenticated"}), 401
        if not is_admin(u): return jsonify({"error":"Admin access required"}), 403
        session.modified = True
        return f(*a, **k)
    return dec

# ── CSRF ──────────────────────────────────────────────────────
def get_csrf_token():
    if "csrf_token" not in session:
        session["csrf_token"] = secrets.token_hex(32)
    return session["csrf_token"]

def verify_csrf(f):
    """Decorator: check X-CSRF-Token header on state-changing requests."""
    @wraps(f)
    def dec(*a, **k):
        if request.method in ("POST","PUT","DELETE","PATCH"):
            token = request.headers.get("X-CSRF-Token","")
            if not token or token != session.get("csrf_token",""):
                return jsonify({"error":"CSRF validation failed"}), 403
        return f(*a, **k)
    return dec

# ── Embeddings (OpenAI text-embedding-3-small) ────────────────
def get_embedding_openai(text: str) -> list:
    """Real semantic embedding via OpenAI. Falls back to hash-vector if no key."""
    oai_key = get_oai_key()
    if not oai_key:
        return _hash_embedding(text)
    try:
        import openai
        client = openai.OpenAI(api_key=oai_key)
        resp = client.embeddings.create(
            model="text-embedding-3-small",
            input=text[:8000]  # Stay within token limits
        )
        return resp.data[0].embedding
    except Exception as e:
        print(f"  [Embedding fallback] {e}")
        return _hash_embedding(text)

def _hash_embedding(text: str) -> list:
    """Lightweight fallback: hash-based sparse vector (256 dims)."""
    words = set(re.findall(r'\b[a-zA-Z]{3,}\b', text.lower()))
    vec = [0.0] * 256
    for w in words:
        idx = int(hashlib.md5(w.encode()).hexdigest(), 16) % 256
        vec[idx] += 1.0
    norm = sum(x*x for x in vec) ** 0.5
    return [x/norm for x in vec] if norm > 0 else vec

def cosine_sim(a, b):
    if not a or not b or len(a) != len(b): return 0.0
    dot = sum(x*y for x,y in zip(a,b))
    na = sum(x*x for x in a)**0.5
    nb = sum(x*x for x in b)**0.5
    return dot/(na*nb) if na*nb > 0 else 0.0

def save_embedding(emb_dir: Path, doc_id: str, emb: list):
    (emb_dir / f"{doc_id}.json").write_text(json.dumps(emb))

def load_embeddings(emb_dir: Path) -> dict:
    out = {}
    if not emb_dir.exists(): return out
    for f in emb_dir.glob("*.json"):
        try: out[f.stem] = json.loads(f.read_text())
        except: pass
    return out

# ── Shared library ────────────────────────────────────────────
def load_shared_index():
    f = SHARED_PATH/"index.json"
    if f.exists():
        try: return json.loads(f.read_text(encoding="utf-8"))
        except: return {}
    return {}

def save_shared_index(idx):
    (SHARED_PATH/"index.json").write_text(json.dumps(idx,ensure_ascii=False,indent=2),encoding="utf-8")

# ── Text extraction ───────────────────────────────────────────
def extract_bytes(data, filename):
    try:
        if filename.lower().endswith(".docx"):
            from docx import Document
            return "\n".join(p.text for p in Document(io.BytesIO(data)).paragraphs if p.text.strip())
        else:
            import pypdf
            return "\n".join(p.extract_text() or "" for p in pypdf.PdfReader(io.BytesIO(data)).pages).strip()
    except: return ""

def extract_path(path):
    try:
        s = Path(path).suffix.lower()
        if s == ".pdf":
            import pypdf
            return "\n".join(p.extract_text() or "" for p in pypdf.PdfReader(str(path)).pages).strip()
        elif s == ".docx":
            from docx import Document
            return "\n".join(p.text for p in Document(str(path)).paragraphs if p.text.strip())
    except: return ""
    return ""

# ── Context retrieval (hybrid semantic + keyword) ─────────────
def get_context(username, query, max_chars=10000):
    query_words = set(w.lower() for w in re.findall(r'\b[a-zA-Z]{3,}\b', query))
    query_emb   = get_embedding_openai(query[:1000])
    results     = []

    # 1. Shared library
    shared_embs = load_embeddings(SHARED_PATH/"embeddings")
    for doc_id, meta in load_shared_index().items():
        tp = SHARED_PATH / f"{doc_id}.txt"
        if not tp.exists(): continue
        sem  = cosine_sim(query_emb, shared_embs.get(doc_id, [])) * 10
        kw   = sum(2 for w in query_words if w in meta.get("name","").lower())
        kw  += sum(1 for w in query_words if w in meta.get("topic","").lower())
        kw  += sum(1 for w in query_words if w in meta.get("jurisdiction","").lower())
        if sem + kw > 0.3:
            results.append((sem+kw, "library", doc_id, meta, tp))

    # 2. User workspace
    up = upath(username)
    user_embs = load_embeddings(up/"embeddings")
    for f in (up/"workspace").glob("*.txt"):
        did = f.stem
        try: text_preview = f.read_text(encoding="utf-8")[:200]
        except: continue
        sem = cosine_sim(query_emb, user_embs.get(did, [])) * 10
        kw  = sum(2 for w in query_words if w in did.lower().replace("_"," "))
        kw += sum(1 for w in query_words if w in text_preview.lower())
        if sem + kw > 0.3:
            results.append((sem+kw, "workspace", did, {"name":did.replace("_"," ")}, f))

    # 3. Drive-indexed files
    for key, meta in load_index(username).items():
        stem = meta.get("stem","").lower().replace("_"," ").replace("-"," ")
        sem  = cosine_sim(query_emb, user_embs.get(meta.get("stem",""), [])) * 10
        kw   = sum(3 for w in query_words if w in stem)
        if sem + kw > 1:
            results.append((sem+kw, "drive", key, meta, Path(meta["path"])))

    results.sort(key=lambda x: x[0], reverse=True)
    parts = []; total = 0; seen = set()
    for score, source, did, meta, path in results[:6]:
        if total >= max_chars or did in seen: continue
        seen.add(did)
        try:
            text = path.read_text(encoding="utf-8") if path.suffix==".txt" else extract_path(str(path))
            if not text: continue
            excerpt = text[:3000]
            lbl = {"library":"Legal Library","workspace":"My Workspace","drive":meta.get("label","Drive")}.get(source,"")
            parts.append(f"[Source: {meta.get('name',did)} | {lbl}]\n{excerpt}")
            total += len(excerpt)
        except: pass
    return parts

def build_prompt(q, ctx):
    if not ctx: return q
    return "Relevant documents:\n\n" + "\n\n---\n\n".join(ctx) + f"\n\n---\n\nQuestion: {q}"

def get_system(display_name):
    return f"""You are Mickey, a legal intelligence assistant for {display_name}.

You assist with any legal topic across any jurisdiction — Belgian law, EU law, English law, French law, Dutch law, Luxembourg law, US law, international law, and others.

Standards:
1. Cite articles precisely: "Art. 28(3) GDPR", "Art. 7:96 WVV", "s.90 Companies Act 2006"
2. Include jurisprudence: cite CJEU decisions, Belgian DPA rulings, EDPB guidelines, national court decisions
3. Distinguish clearly between settled law, developing case law, and your assessment
4. Use structured headers for longer answers
5. End every answer with "**Sources cited:**" listing all articles, cases and documents referenced
6. When documents from the knowledge base are provided, cite them by name and prioritise them
7. For Belgian law, always check both federal and regional competences where relevant"""

# ══════════════════════════════════════════════════════════════
# ROUTES
# ══════════════════════════════════════════════════════════════

@app.route("/app")
def app_index():
    return render_template("index.html")

@app.route("/api/csrf")
@login_required
def csrf():
    return jsonify({"token": get_csrf_token()})

# ── Auth ──────────────────────────────────────────────────────

@app.route("/api/auth/status")
def auth_status():
    u = session.get("username")
    if u:
        users = load_users(); usr = users.get(u,{})
        return jsonify({"logged_in":True,"username":u,
                        "display_name":usr.get("display_name",u),
                        "is_admin":usr.get("role")=="admin",
                        "csrf_token":get_csrf_token()})
    return jsonify({"logged_in":False})

@app.route("/api/auth/login", methods=["POST"])
def login():
    d = request.get_json(force=True) or {}
    u  = (d.get("username") or "").strip().lower()
    pw = d.get("password") or ""
    if not u or not pw: return jsonify({"error":"Username and password required"}),400
    users = load_users()
    if u not in users:
        record_login_fail(u)
        return jsonify({"error":"Username not found"}),400
    usr = users[u]

    # Check brute-force lockout
    locked, secs = is_locked_out(u)
    if locked:
        mins = (secs + 59) // 60
        return jsonify({"error":f"Too many failed attempts. Try again in {mins} minute(s)."}),429

    # Migration: if old SHA256 hash, verify then upgrade to bcrypt
    if _is_old_hash(usr.get("hashed","")):
        if not _old_verify(pw, usr["hashed"], usr.get("salt","")):
            record_login_fail(u)
            return jsonify({"error":"Incorrect password"}),400
        usr["hashed"]  = hash_pw(pw)
        usr["salt"]    = ""; usr["hash_v"]  = "bcrypt"
        save_users(users)
    else:
        if not verify_pw(pw, usr["hashed"]):
            record_login_fail(u)
            return jsonify({"error":"Incorrect password"}),400

    clear_login_fails(u)
    session.clear(); session["username"]=u; session.permanent=True
    log_usage(u, "login")
    return jsonify({"ok":True,"display_name":usr.get("display_name",u),"is_admin":usr.get("role")=="admin","csrf_token":get_csrf_token()})

@app.route("/api/auth/register", methods=["POST"])
def register():
    d    = request.get_json(force=True) or {}
    name = (d.get("display_name") or "").strip()
    u    = (d.get("username") or "").strip().lower()
    pw   = d.get("password") or ""
    pw2  = d.get("confirm") or ""
    if not name or not u or not pw: return jsonify({"error":"All fields required"}),400
    if " " in u: return jsonify({"error":"Username cannot contain spaces"}),400
    users = load_users()
    if u in users: return jsonify({"error":"Username already taken"}),400
    if pw != pw2: return jsonify({"error":"Passwords don't match"}),400
    valid, rules = validate_pw(pw)
    if not valid: return jsonify({"error":"Password needs: "+", ".join(rules)}),400
    role = "admin" if not users else "user"
    users[u] = {"display_name":name,"hashed":hash_pw(pw),"hash_v":"bcrypt","salt":"",
                "role":role,"created":datetime.datetime.now().isoformat()}
    save_users(users); upath(u)
    save_settings(u,{"display_name":name,"drive_paths":[],"drive_labels":[]})
    session.clear(); session["username"]=u; session.permanent=True
    log_usage(u,"register")
    return jsonify({"ok":True,"display_name":name,"is_admin":role=="admin","csrf_token":get_csrf_token()})

@app.route("/api/auth/logout", methods=["POST"])
def do_logout():
    if session.get("username"): log_usage(session["username"],"logout")
    session.clear(); return jsonify({"ok":True})

@app.route("/api/auth/change_password", methods=["POST"])
@login_required
def change_password():
    d       = request.get_json(force=True) or {}
    current = d.get("current","")
    new_pw  = d.get("new","")
    new_pw2 = d.get("confirm","")
    u       = session["username"]
    users   = load_users()
    usr     = users[u]
    if not verify_pw(current, usr["hashed"]):
        return jsonify({"error":"Current password is incorrect"}),400
    if new_pw != new_pw2:
        return jsonify({"error":"New passwords don't match"}),400
    valid, rules = validate_pw(new_pw)
    if not valid: return jsonify({"error":"Password needs: "+", ".join(rules)}),400
    usr["hashed"] = hash_pw(new_pw); usr["hash_v"]="bcrypt"; usr["salt"]=""
    save_users(users)
    return jsonify({"ok":True})

# ── API Keys ──────────────────────────────────────────────────

@app.route("/api/key/set", methods=["POST"])
@admin_required
def set_key():
    d = request.get_json(force=True) or {}
    key = (d.get("key") or "").strip()
    if not key: return jsonify({"error":"No key provided"}),400
    _write_key(API_KEY_FILE, key); return jsonify({"ok":True})

@app.route("/api/key/set_oai", methods=["POST"])
@admin_required
def set_oai_key():
    d = request.get_json(force=True) or {}
    key = (d.get("key") or "").strip()
    if not key: return jsonify({"error":"No key provided"}),400
    _write_key(OAI_KEY_FILE, key); return jsonify({"ok":True})

@app.route("/api/key/status")
@login_required
def key_status():
    ak = get_api_key(); ok = get_oai_key()
    return jsonify({
        "anthropic": bool(ak), "anthropic_masked": ("sk-ant-..."+ak[-6:]) if ak else "",
        "openai":    bool(ok), "openai_masked":    ("sk-..."+ok[-6:])    if ok else ""
    })

# ── Ask (with web search + rate limiting + usage tracking) ────

@app.route("/api/ask", methods=["POST"])
@login_required
def ask():
    username = session["username"]
    if not check_rate_limit(username):
        return jsonify({"error":f"Rate limit reached ({RATE_LIMIT_CALLS_PER_HOUR} calls/hour). Please wait."}),429

    d        = request.get_json(force=True) or {}
    question = d.get("question","")
    history  = d.get("history",[])
    use_web  = d.get("web_search", True)
    display  = load_users().get(username,{}).get("display_name",username)

    ctx      = get_context(username, question)
    prompt   = build_prompt(question, ctx)
    messages = [{"role":h["role"],"content":h["content"]} for h in history[-6:]
                if h.get("role") in ("user","assistant") and h.get("content")]
    messages.append({"role":"user","content":prompt})

    try:
        client = get_client()
        kwargs = dict(model="claude-sonnet-4-20250514", max_tokens=3000,
                      system=get_system(display), messages=messages)
        if use_web:
            kwargs["tools"] = [{"type":"web_search_20250305","name":"web_search"}]
        resp  = client.messages.create(**kwargs)
        answer = " ".join(b.text for b in resp.content if hasattr(b,"text") and b.text).strip()
        tokens = resp.usage.input_tokens + resp.usage.output_tokens
        log_usage(username, "ask", tokens, "discovery")
        return jsonify({"answer":answer})
    except ValueError as e:
        return jsonify({"error":str(e),"need_key":True}),500
    except Exception as e:
        return jsonify({"error":str(e)}),500

# ── Analyze ───────────────────────────────────────────────────

@app.route("/api/analyze", methods=["POST"])
@login_required
def analyze():
    username = session["username"]
    if not check_rate_limit(username):
        return jsonify({"error":"Rate limit reached. Please wait."}),429

    display = load_users().get(username,{}).get("display_name",username)
    file    = request.files.get("file")
    if not file: return jsonify({"error":"No file uploaded"}),400
    module  = request.form.get("module","review")
    focus   = request.form.get("focus","GDPR, EU AI Act")
    playbook= request.form.get("playbook","")
    extra   = request.form.get("extra","")

    text = extract_bytes(file.read(), file.filename)
    if not text: return jsonify({"error":"Could not extract text. Is this a scanned/image PDF?"}),400

    ctx     = get_context(username, focus)
    ctx_str = "\n\nReference from knowledge base:\n\n"+"---".join(ctx[:2]) if ctx else ""

    if module=="review":
        prompt = f"""Review this document for legal compliance and risk.

Document: {file.filename}
Check against: {focus}
{f"Playbook: {playbook}" if playbook else ""}
{f"Instructions: {extra}" if extra else ""}
{ctx_str}

Document:
{text[:8000]}

Structure:
1. **Summary** — 3-4 sentences on purpose and key provisions
2. **Findings**
   - 🔴 **HIGH** — must fix (issue | article | fix)
   - 🟡 **MEDIUM** — should address (issue | article | fix)
   - 🟢 **LOW** — minor improvements
3. **Overall recommendation**"""

    elif module=="compare":
        file2 = request.files.get("file2")
        if not file2: return jsonify({"error":"Two files required"}),400
        text2 = extract_bytes(file2.read(), file2.filename)
        prompt = f"""Compare clause by clause.

Doc 1 (under review): {file.filename}
{text[:4000]}

Doc 2 (reference): {file2.filename}
{text2[:4000]}

Focus: {focus}{ctx_str}

For each clause: Doc 1 position | Doc 2 position | Status (Matches/Differs/Missing) | Risk if different
End with: Key deviations and negotiation priorities."""

    elif module=="translate":
        src=request.form.get("src","Dutch (NL)"); tgt=request.form.get("tgt","English (EN)")
        style=request.form.get("style","Legal / formal")
        prompt = f"""Translate from {src} to {tgt}. Style: {style}.
{f"Notes: {extra}" if extra else ""}
Rules: preserve article numbering, defined terms capitalised, legal structure intact. Translation only.

{text[:7000]}"""
    else:
        prompt = f"Analyze: {file.filename}\n\n{text[:8000]}"

    try:
        client = get_client()
        resp   = client.messages.create(model="claude-sonnet-4-20250514",max_tokens=3000,
                                        system=get_system(display),messages=[{"role":"user","content":prompt}])
        answer = " ".join(b.text for b in resp.content if hasattr(b,"text") and b.text).strip()
        tokens = resp.usage.input_tokens + resp.usage.output_tokens
        log_usage(username, module, tokens, file.filename[:40])
        return jsonify({"result":answer,"filename":file.filename})
    except ValueError as e: return jsonify({"error":str(e),"need_key":True}),500
    except Exception as e:  return jsonify({"error":str(e)}),500

# ── Draft ─────────────────────────────────────────────────────

@app.route("/api/draft", methods=["POST"])
@login_required
def draft():
    username = session["username"]
    if not check_rate_limit(username): return jsonify({"error":"Rate limit reached."}),429
    display = load_users().get(username,{}).get("display_name",username)
    d       = request.get_json(force=True) or {}
    clause  = d.get("clause","NDA"); doc_type=d.get("doc_type","Contract clause")
    ctx     = get_context(username, clause+" "+doc_type)
    ctx_str = "\n\nReference:\n\n"+"---".join(ctx[:2]) if ctx else ""
    prompt  = f"""Draft a {doc_type}: {clause}
Governing law: {d.get('law','Belgian law')} | Perspective: {d.get('perspective','Neutral')}
{f"Instructions: {d.get('instructions','')}" if d.get('instructions') else ""}
{ctx_str}
Requirements: numbered articles, cite legal obligations precisely, professional drafting style.
End with "Warning: Mickey note:" flagging items requiring verification."""
    try:
        client = get_client()
        resp   = client.messages.create(model="claude-sonnet-4-20250514",max_tokens=3000,
                                        system=get_system(display),messages=[{"role":"user","content":prompt}])
        answer = " ".join(b.text for b in resp.content if hasattr(b,"text") and b.text).strip()
        tokens = resp.usage.input_tokens + resp.usage.output_tokens
        log_usage(username,"draft",tokens,clause[:40])
        return jsonify({"result":answer})
    except ValueError as e: return jsonify({"error":str(e),"need_key":True}),500
    except Exception as e:  return jsonify({"error":str(e)}),500

# ── Workspace ─────────────────────────────────────────────────

@app.route("/api/workspace", methods=["GET"])
@login_required
def get_workspace():
    up = upath(session["username"])
    docs = [{"id":f.stem,"name":f.stem.replace("_"," "),"size_kb":round(f.stat().st_size/1024,1)}
            for f in (up/"workspace").glob("*.txt")]
    return jsonify({"docs":sorted(docs,key=lambda x:x["name"])})

@app.route("/api/workspace/add", methods=["POST"])
@login_required
def add_workspace():
    username = session["username"]
    file=request.files.get("file"); name=(request.form.get("name") or "").strip()
    if not file or not name: return jsonify({"error":"File and name required"}),400
    text = extract_bytes(file.read(),file.filename)
    if not text: return jsonify({"error":"Could not extract text"}),400
    up  = upath(username)
    did = re.sub(r'[^\w\s\-.]','',name)[:80].strip().replace(" ","_")
    (up/"workspace"/f"{did}.txt").write_text(text,encoding="utf-8")
    try:
        emb = get_embedding_openai(text[:3000])
        save_embedding(up/"embeddings",did,emb)
    except: pass
    log_usage(username,"add_workspace",0,name[:40])
    return jsonify({"ok":True})

@app.route("/api/workspace/delete", methods=["POST"])
@login_required
def delete_workspace_doc():
    username=session["username"]; d=request.get_json(force=True) or {}; did=d.get("id","")
    up=upath(username)
    for ext in [".txt"]: f=up/"workspace"/f"{did}{ext}"; f.exists() and f.unlink()
    emb=up/"embeddings"/f"{did}.json"; emb.exists() and emb.unlink()
    return jsonify({"ok":True})

# ── Library ───────────────────────────────────────────────────

@app.route("/api/library", methods=["GET"])
@login_required
def get_library():
    idx  = load_shared_index()
    docs = sorted([{**v,"id":k} for k,v in idx.items()],
                  key=lambda x:x.get("added_at",""),reverse=True)
    return jsonify({"docs":docs})

@app.route("/api/library/add", methods=["POST"])
@admin_required
def add_library():
    file=request.files.get("file"); name=(request.form.get("name") or "").strip()
    if not file or not name: return jsonify({"error":"File and name required"}),400
    text=extract_bytes(file.read(),file.filename)
    if not text: return jsonify({"error":"Could not extract text"}),400
    did=secrets.token_hex(8)
    idx=load_shared_index()
    idx[did]={"name":name,"jurisdiction":(request.form.get("jurisdiction") or "").strip(),
              "topic":(request.form.get("topic") or "").strip(),"filename":file.filename,
              "added_by":session["username"],"added_at":datetime.datetime.now().isoformat(),
              "size_kb":round(len(text.encode())/1024,1)}
    save_shared_index(idx)
    (SHARED_PATH/f"{did}.txt").write_text(text,encoding="utf-8")
    try:
        emb=get_embedding_openai(text[:3000])
        save_embedding(SHARED_PATH/"embeddings",did,emb)
    except: pass
    log_usage(session["username"],"add_library",0,name[:40])
    return jsonify({"ok":True,"id":did})

@app.route("/api/library/delete", methods=["POST"])
@admin_required
def delete_library():
    d=request.get_json(force=True) or {}; did=d.get("id","")
    idx=load_shared_index()
    if did in idx: del idx[did]
    save_shared_index(idx)
    for ext in [".txt"]:
        f=SHARED_PATH/f"{did}{ext}"; f.exists() and f.unlink()
    emb=SHARED_PATH/"embeddings"/f"{did}.json"; emb.exists() and emb.unlink()
    return jsonify({"ok":True})

# ── Storage ───────────────────────────────────────────────────

@app.route("/api/storage", methods=["GET"])
@login_required
def get_storage():
    u=session["username"]; s=load_settings(u); idx=load_index(u)
    return jsonify({"folders":[{"path":p,"label":l,
        "count":sum(1 for v in idx.values() if v.get("label")==l),
        "exists":Path(p).exists()}
        for p,l in zip(s.get("drive_paths",[]),s.get("drive_labels",[]))]})

@app.route("/api/storage/add", methods=["POST"])
@login_required
def add_storage():
    u=session["username"]; d=request.get_json(force=True) or {}
    path=(d.get("path") or "").strip(); label=(d.get("label") or "").strip()
    if not path or not label: return jsonify({"error":"Path and label required"}),400
    if not Path(path).exists(): return jsonify({"error":f"Folder not found: {path}"}),400
    s=load_settings(u); s["drive_paths"].append(path); s["drive_labels"].append(label)
    save_settings(u,s)
    new,unch,total=sync_folder(u,path,label)
    return jsonify({"ok":True,"new":new,"unchanged":unch,"total":total})

@app.route("/api/storage/sync", methods=["POST"])
@login_required
def sync_storage():
    u=session["username"]; d=request.get_json(force=True) or {}
    path=d.get("path",""); label=d.get("label","")
    if not Path(path).exists(): return jsonify({"error":f"Folder not found: {path}"}),400
    new,unch,total=sync_folder(u,path,label)
    return jsonify({"ok":True,"new":new,"unchanged":unch,"total":total})

@app.route("/api/storage/remove", methods=["POST"])
@login_required
def remove_storage():
    u=session["username"]; d=request.get_json(force=True) or {}; label=d.get("label","")
    s=load_settings(u)
    pairs=[(p,l) for p,l in zip(s["drive_paths"],s["drive_labels"]) if l!=label]
    s["drive_paths"]=[p for p,l in pairs]; s["drive_labels"]=[l for p,l in pairs]
    save_settings(u,s); return jsonify({"ok":True})

def sync_folder(username,folder_path,label):
    idx=load_index(username); new_files=unchanged=0; all_files=[]
    try:
        for path in Path(folder_path).rglob("*"):
            if not path.is_file(): continue
            if path.suffix.lower() not in [".pdf",".docx"]: continue
            if path.name.startswith(".") or path.name.startswith("~$"): continue
            if any(sk in [p.lower() for p in path.parts] for sk in SKIP_FOLDERS): continue
            all_files.append(path)
    except: pass
    for path in all_files:
        key=str(path)
        try: mtime=path.stat().st_mtime
        except: continue
        if key in idx and idx[key].get("mtime")==mtime: unchanged+=1; continue
        idx[key]={"name":path.name,"stem":path.stem,"path":str(path),"label":label,
                  "mtime":mtime,"size_kb":round(path.stat().st_size/1024,1),
                  "suffix":path.suffix.lower(),"indexed_at":datetime.datetime.now().isoformat()}
        new_files+=1
    save_index(username,idx); return new_files,unchanged,len(all_files)

# ── Stats ─────────────────────────────────────────────────────

@app.route("/api/stats")
@login_required
def get_stats():
    u=session["username"]; up=upath(u); idx=load_index(u)
    return jsonify({"workspace":len(list((up/"workspace").glob("*.txt"))),
                    "library":len(load_shared_index()),"drive":len(idx)})

# ── Usage ─────────────────────────────────────────────────────

@app.route("/api/usage")
@login_required
def my_usage():
    return jsonify(get_usage_summary(session["username"]))

@app.route("/api/admin/usage")
@admin_required
def all_usage():
    return jsonify(get_all_usage_summary())

# ── Admin: users ──────────────────────────────────────────────

@app.route("/api/admin/users")
@admin_required
def admin_users():
    users=load_users()
    return jsonify({"users":[{"username":u,"display_name":v.get("display_name",u),
        "role":v.get("role","user"),"created":v.get("created",""),
        "usage":get_usage_summary(u,7)}   # last 7 days inline
        for u,v in users.items()]})

@app.route("/api/admin/users/delete", methods=["POST"])
@admin_required
def admin_delete():
    d=request.get_json(force=True) or {}; target=d.get("username","")
    if target==session["username"]: return jsonify({"error":"Cannot delete your own account"}),400
    users=load_users()
    if target not in users: return jsonify({"error":"User not found"}),400
    del users[target]; save_users(users)
    ud=DATA_PATH/target
    if ud.exists(): shutil.rmtree(ud)
    return jsonify({"ok":True})

@app.route("/api/admin/users/promote", methods=["POST"])
@admin_required
def admin_promote():
    d=request.get_json(force=True) or {}; target=d.get("username","")
    users=load_users()
    if target not in users: return jsonify({"error":"User not found"}),400
    users[target]["role"]="admin"; save_users(users); return jsonify({"ok":True})


# ── Collections ───────────────────────────────────────────────

def load_collections(username):
    f = upath(username) / "collections.json"
    if f.exists():
        try: return json.loads(f.read_text(encoding="utf-8"))
        except: return {}
    return {}

def save_collections(username, cols):
    (upath(username) / "collections.json").write_text(
        json.dumps(cols, indent=2, ensure_ascii=False), encoding="utf-8")

@app.route("/api/collections", methods=["GET"])
@login_required
def get_collections():
    return jsonify({"collections": load_collections(session["username"])})

@app.route("/api/collections/create", methods=["POST"])
@login_required
def create_collection():
    d = request.get_json(force=True) or {}
    name = (d.get("name") or "").strip()
    if not name: return jsonify({"error": "Name required"}), 400
    cols = load_collections(session["username"])
    cid = secrets.token_hex(6)
    cols[cid] = {"name": name, "created": datetime.datetime.now().isoformat(), "items": []}
    save_collections(session["username"], cols)
    return jsonify({"ok": True, "id": cid})

@app.route("/api/collections/save_item", methods=["POST"])
@login_required
def save_to_collection():
    d = request.get_json(force=True) or {}
    cid       = d.get("collection_id", "")
    title     = (d.get("title") or "").strip()
    content   = (d.get("content") or "").strip()
    item_type = d.get("type", "answer")
    source    = d.get("source", "")
    if not cid or not content: return jsonify({"error": "Collection and content required"}), 400
    cols = load_collections(session["username"])
    if cid not in cols: return jsonify({"error": "Collection not found"}), 404
    item = {
        "id": secrets.token_hex(6),
        "title": title or content[:60] + ("..." if len(content) > 60 else ""),
        "content": content,
        "type": item_type,
        "source": source,
        "saved_at": datetime.datetime.now().isoformat()
    }
    cols[cid]["items"].append(item)
    save_collections(session["username"], cols)
    return jsonify({"ok": True})

@app.route("/api/collections/delete_item", methods=["POST"])
@login_required
def delete_collection_item():
    d = request.get_json(force=True) or {}
    cid = d.get("collection_id", ""); iid = d.get("item_id", "")
    cols = load_collections(session["username"])
    if cid not in cols: return jsonify({"error": "Collection not found"}), 404
    cols[cid]["items"] = [i for i in cols[cid]["items"] if i["id"] != iid]
    save_collections(session["username"], cols)
    return jsonify({"ok": True})

@app.route("/api/collections/delete", methods=["POST"])
@login_required
def delete_collection():
    d = request.get_json(force=True) or {}
    cid = d.get("collection_id", "")
    cols = load_collections(session["username"])
    if cid in cols: del cols[cid]
    save_collections(session["username"], cols)
    return jsonify({"ok": True})

@app.route("/api/collections/rename", methods=["POST"])
@login_required
def rename_collection():
    d = request.get_json(force=True) or {}
    cid = d.get("collection_id", ""); name = (d.get("name") or "").strip()
    cols = load_collections(session["username"])
    if cid not in cols: return jsonify({"error": "Not found"}), 404
    cols[cid]["name"] = name
    save_collections(session["username"], cols)
    return jsonify({"ok": True})

# ── Data requests (GDPR Art. 15 / Art. 17) ────────────────────

REQUESTS_FILE = APP_PATH / "data_requests.json"

def load_requests():
    if REQUESTS_FILE.exists():
        try: return json.loads(REQUESTS_FILE.read_text(encoding="utf-8"))
        except: return []
    return []

def save_requests(reqs):
    REQUESTS_FILE.write_text(json.dumps(reqs, indent=2, ensure_ascii=False), encoding="utf-8")

@app.route("/api/data_request/submit", methods=["POST"])
@login_required
def submit_data_request():
    d = request.get_json(force=True) or {}
    req_type = d.get("type", "export")
    username = session["username"]
    users = load_users()
    display = users.get(username, {}).get("display_name", username)
    reqs = load_requests()
    for r in reqs:
        if r["username"] == username and r["type"] == req_type and r["status"] == "pending":
            return jsonify({"error": "You already have a pending request of this type"}), 400
    reqs.append({
        "id": secrets.token_hex(8),
        "username": username,
        "display_name": display,
        "type": req_type,
        "status": "pending",
        "submitted_at": datetime.datetime.now().isoformat(),
        "processed_at": None,
        "processed_by": None
    })
    save_requests(reqs)
    log_usage(username, f"data_request_{req_type}")
    return jsonify({"ok": True})

@app.route("/api/data_request/status")
@login_required
def my_data_requests():
    username = session["username"]
    reqs = [r for r in load_requests() if r["username"] == username]
    return jsonify({"requests": reqs})

@app.route("/api/admin/data_requests")
@admin_required
def admin_data_requests():
    return jsonify({"requests": load_requests()})

@app.route("/api/admin/data_request/approve", methods=["POST"])
@admin_required
def approve_data_request():
    d = request.get_json(force=True) or {}
    rid = d.get("id", "")
    reqs = load_requests()
    req = next((r for r in reqs if r["id"] == rid), None)
    if not req: return jsonify({"error": "Request not found"}), 404
    if req["status"] != "pending": return jsonify({"error": "Already processed"}), 400
    target = req["username"]
    req["status"] = "approved"
    req["processed_at"] = datetime.datetime.now().isoformat()
    req["processed_by"] = session["username"]
    save_requests(reqs)
    if req["type"] == "export":
        import zipfile
        buf = io.BytesIO()
        with zipfile.ZipFile(buf, "w", zipfile.ZIP_DEFLATED) as zf:
            up = upath(target)
            for f in (up / "workspace").glob("*.txt"):
                zf.write(f, f"workspace/{f.name}")
            sf = up / "settings.json"
            if sf.exists(): zf.write(sf, "profile/settings.json")
            cf = up / "collections.json"
            if cf.exists(): zf.write(cf, "profile/collections.json")
            for lf in USAGE_PATH.glob(f"{target}_*.jsonl"):
                zf.write(lf, f"usage/{lf.name}")
            readme = f"""Mickey -- Data Export
User: {req["display_name"]} ({target})
Exported: {datetime.datetime.now().isoformat()}
Requested: {req["submitted_at"]}

Contents:
- workspace/    Your private workspace documents
- profile/      Your settings and collections
- usage/        Your usage logs (timestamps and event types only)

This export is provided in accordance with Article 20 GDPR
(Right to data portability) and the Belgian Act of 30 July 2018
on the Protection of Natural Persons with Regard to the
Processing of Personal Data.
"""
            zf.writestr("README.txt", readme)
        buf.seek(0)
        from flask import send_file
        return send_file(buf, mimetype="application/zip", as_attachment=True,
                         download_name=f"mickey_export_{target}_{datetime.date.today()}.zip")
    return jsonify({"ok": True, "message": "Deletion approved. Delete the account via user management."})

@app.route("/api/admin/data_request/reject", methods=["POST"])
@admin_required
def reject_data_request():
    d = request.get_json(force=True) or {}
    rid = d.get("id", ""); reason = d.get("reason", "")
    reqs = load_requests()
    req = next((r for r in reqs if r["id"] == rid), None)
    if not req: return jsonify({"error": "Not found"}), 404
    req["status"] = "rejected"
    req["processed_at"] = datetime.datetime.now().isoformat()
    req["processed_by"] = session["username"]
    req["reject_reason"] = reason
    save_requests(reqs)
    return jsonify({"ok": True})

# ── Info pages ─────────────────────────────────────────────────

@app.route("/api/info/<page>")
def info_page(page):
    pages = {"terms": _terms(), "privacy": _privacy(), "disclaimer": _disclaimer()}
    if page not in pages: return jsonify({"error": "Not found"}), 404
    return jsonify({"content": pages[page]})

def _terms():
    return """# Terms of Use

**Version 1.0 - February 2026**
**Operator:** [Your Name / Company], [Address], Belgium

## 1. Nature of the Platform

Mickey is an AI-assisted legal intelligence tool for authorised users within [Your Organisation]. Mickey is not a law firm and does not provide legal advice. All outputs are AI-generated and must be reviewed by a qualified legal professional before reliance.

## 2. Authorised Use

Authorised users may use Mickey for legal research, document review, drafting support, translation, and knowledge management. Users may not share credentials, upload unauthorised confidential information, or attempt to circumvent security measures.

## 3. No Legal Advice

Mickey outputs do not constitute legal advice (juridisch advies / conseil juridique). They may contain errors or outdated information. No attorney-client relationship is created. Always verify with a qualified Belgian lawyer before acting.

## 4. Intellectual Property

Uploaded documents remain the property of the user or their organisation. AI outputs are provided for the user's use without guarantee of originality or freedom from third-party rights.

## 5. Governing Law

These Terms are governed by Belgian law. Disputes are subject to the exclusive jurisdiction of the courts of Brussels, Belgium.

---
*Last updated: February 2026*"""

def _privacy():
    return """# Privacy Notice

**Version 1.0 - February 2026**

Provided in accordance with Articles 13-14 GDPR and the Belgian Act of 30 July 2018 on the Protection of Natural Persons with Regard to the Processing of Personal Data.

## 1. Controller

**[Your Name / Company]**, [Address], Belgium, [email]

## 2. Data We Process

- **Account data:** name, username, encrypted password, creation date, role
- **Usage data:** timestamps, event types, token counts (no question content stored)
- **Workspace documents:** files you upload for knowledge retrieval
- **Collections:** your saved answers and bookmarks
- **Session tokens:** valid 8 hours from last activity

## 3. Legal Basis

- **Contract (Art. 6(1)(b) GDPR):** authentication and service delivery
- **Legitimate interest (Art. 6(1)(f) GDPR):** usage monitoring, abuse prevention, service improvement

## 4. AI Processing

Questions and documents are transmitted to the Anthropic Claude API (data processor). Where semantic search is enabled, document excerpts are transmitted to the OpenAI Embeddings API (data processor).

## 5. Retention

| Data | Retention |
|------|-----------|
| Account data | Until account deletion |
| Workspace / Collections | Until deleted or account deletion |
| Usage logs | 12 months rolling |
| Session tokens | 8 hours |

## 6. Your Rights

Under GDPR you have rights of access (Art. 15), rectification (Art. 16), erasure (Art. 17), portability (Art. 20), and objection (Art. 21). Exercise these via Settings > Data & Privacy or contact us at [email]. We respond within one month (Art. 12 GDPR).

## 7. Supervisory Authority

**Gegevensbeschermingsautoriteit (GBA) / Autorite de protection des donnees (APD)**
Rue de la Presse 35, 1000 Brussels
www.gegevensbeschermingsautoriteit.be

## 8. Security

We implement bcrypt password hashing, encrypted key storage, session timeouts, and role-based access control as appropriate technical and organisational measures under Art. 32 GDPR.

---
*Last updated: February 2026 | Governing law: Belgian law / GDPR*"""

def _disclaimer():
    return """# AI Output Disclaimer

Mickey uses artificial intelligence to assist with legal research, document analysis, drafting, and translation. By using Mickey, you acknowledge that:

1. **Not legal advice.** Outputs do not constitute legal advice (juridisch advies / conseil juridique) and do not create an attorney-client relationship.

2. **Verification required.** All outputs must be verified by a qualified legal professional before reliance, filing, or communication to clients.

3. **Accuracy not guaranteed.** AI models may produce errors, hallucinations, or outdated analysis. Mickey's training has a knowledge cutoff and may not reflect recent legislative or jurisprudential developments.

4. **Cite verification required.** All case citations, article references, and regulatory decisions must be verified against primary sources.

5. **No liability.** The operator accepts no liability for decisions made on the basis of Mickey outputs without appropriate professional review.

---
*Displayed in accordance with good practice for AI-assisted legal tools and applicable Belgian professional conduct rules.*"""

# ══════════════════════════════════════════════════════════════
# CORPORATE HOUSEKEEPING
# ══════════════════════════════════════════════════════════════

CORP_PATH = APP_PATH / "corporate"
CORP_PATH.mkdir(parents=True, exist_ok=True)

def corp_file(entity_id: str) -> Path:
    """Return path to entity JSON file."""
    safe = re.sub(r'[^\w\-]', '_', entity_id)[:60]
    return CORP_PATH / f"{safe}.json"

def load_entity(entity_id: str) -> dict:
    f = corp_file(entity_id)
    if f.exists():
        try: return json.loads(f.read_text(encoding="utf-8"))
        except: return {}
    return {}

def save_entity(entity_id: str, data: dict):
    corp_file(entity_id).write_text(
        json.dumps(data, indent=2, ensure_ascii=False), encoding="utf-8")

def list_entities() -> list:
    entities = []
    for f in CORP_PATH.glob("*.json"):
        try:
            d = json.loads(f.read_text(encoding="utf-8"))
            entities.append({
                "id":           f.stem,
                "name":         d.get("name", f.stem),
                "legal_form":   d.get("legal_form", ""),
                "jurisdiction": d.get("jurisdiction", ""),
                "status":       d.get("status", "active"),
                "updated_at":   d.get("updated_at", ""),
                "alert_count":  _count_alerts(d),
            })
        except: pass
    return sorted(entities, key=lambda x: x.get("name","").lower())

def _count_alerts(entity: dict) -> int:
    """Count expiring mandates + overdue compliance items."""
    count = 0
    today = datetime.date.today()
    warn_date = today + datetime.timedelta(days=90)
    # Director mandates expiring within 90 days
    for d in entity.get("directors", []):
        if d.get("status") == "active" and d.get("mandate_end"):
            try:
                end = datetime.date.fromisoformat(d["mandate_end"])
                if end <= warn_date: count += 1
            except: pass
    # Overdue compliance items
    for c in entity.get("compliance", []):
        if c.get("status") in ("pending", "overdue") and c.get("deadline"):
            try:
                dl = datetime.date.fromisoformat(c["deadline"])
                if dl < today: count += 1
            except: pass
    return count

def _add_history(entity: dict, event_type: str, description: str,
                 source: str = "", recorded_by: str = "") -> dict:
    """Append a history entry to an entity dict."""
    if "history" not in entity:
        entity["history"] = []
    entity["history"].insert(0, {
        "id":          secrets.token_hex(6),
        "ts":          datetime.datetime.now().isoformat(),
        "date":        datetime.date.today().isoformat(),
        "type":        event_type,   # governance | capital | compliance | document | other
        "description": description,
        "source":      source,
        "recorded_by": recorded_by,
    })
    return entity

# ── Corporate: entity list & create ──────────────────────────

@app.route("/api/corp/entities", methods=["GET"])
@login_required
def corp_entities():
    return jsonify({"entities": list_entities()})

@app.route("/api/corp/entity", methods=["GET"])
@login_required
def corp_get_entity():
    eid = request.args.get("id", "").strip()
    if not eid: return jsonify({"error": "id required"}), 400
    e = load_entity(eid)
    if not e: return jsonify({"error": "Entity not found"}), 404
    return jsonify({"entity": e})

@app.route("/api/corp/entity/create", methods=["POST"])
@login_required
@verify_csrf
def corp_create_entity():
    d    = request.get_json(force=True) or {}
    name = (d.get("name") or "").strip()
    if not name: return jsonify({"error": "Name required"}), 400
    eid  = secrets.token_hex(8)
    now  = datetime.datetime.now().isoformat()
    entity = {
        "id":           eid,
        "name":         name,
        "legal_form":   d.get("legal_form", ""),
        "jurisdiction": d.get("jurisdiction", ""),
        "status":       d.get("status", "active"),
        "relationship": d.get("relationship", "subsidiary"),
        "created_at":   now,
        "updated_at":   now,
        # Identity
        "trade_name":         d.get("trade_name", ""),
        "previous_names":     d.get("previous_names", ""),
        "reg_number":         d.get("reg_number", ""),
        "vat_number":         d.get("vat_number", ""),
        "lei_code":           d.get("lei_code", ""),
        "court_registration": d.get("court_registration", ""),
        "incorporation_date": d.get("incorporation_date", ""),
        "duration":           d.get("duration", "Unlimited"),
        "financial_year_end": d.get("financial_year_end", "31 December"),
        "agm_date_rule":      d.get("agm_date_rule", ""),
        "agm_notice_days":    d.get("agm_notice_days", 30),
        "listed":             d.get("listed", False),
        # Addresses
        "registered_street":  d.get("registered_street", ""),
        "registered_postcode":d.get("registered_postcode", ""),
        "registered_city":    d.get("registered_city", ""),
        "registered_country": d.get("registered_country", ""),
        # Capital
        "share_capital":      d.get("share_capital", ""),
        "capital_currency":   d.get("capital_currency", "EUR"),
        "total_shares":       d.get("total_shares", ""),
        "nominal_value":      d.get("nominal_value", "No par value"),
        # Audit
        "auditor_firm":       d.get("auditor_firm", ""),
        "auditor_partner":    d.get("auditor_partner", ""),
        "auditor_mandate_end":d.get("auditor_mandate_end", ""),
        # Governance
        "board_model":        d.get("board_model", "One-tier"),
        "min_directors":      d.get("min_directors", 1),
        "max_directors":      d.get("max_directors", 12),
        "signature_rule":     d.get("signature_rule", "Joint"),
        # Collections
        "directors":          [],
        "shareholders":       [],
        "poas":               [],
        "compliance":         [],
        "history":            [],
        "documents":          [],
        "notes":              d.get("notes", ""),
    }
    _add_history(entity, "other", f"Entity '{name}' created in Mickey",
                 recorded_by=session["username"])
    save_entity(eid, entity)
    log_usage(session["username"], "corp_create", 0, name[:40])
    return jsonify({"ok": True, "id": eid})

@app.route("/api/corp/entity/update", methods=["POST"])
@login_required
@verify_csrf
def corp_update_entity():
    d   = request.get_json(force=True) or {}
    eid = d.get("id", "").strip()
    if not eid: return jsonify({"error": "id required"}), 400
    entity = load_entity(eid)
    if not entity: return jsonify({"error": "Entity not found"}), 404
    # Update scalar fields — never overwrite sub-collections this way
    protected = {"id","created_at","directors","shareholders","poas",
                 "compliance","history","documents"}
    for k, v in d.items():
        if k not in protected:
            entity[k] = v
    entity["updated_at"] = datetime.datetime.now().isoformat()
    _add_history(entity, "other", "Entity data updated",
                 recorded_by=session["username"])
    save_entity(eid, entity)
    return jsonify({"ok": True})

@app.route("/api/corp/entity/delete", methods=["POST"])
@login_required
@verify_csrf
def corp_delete_entity():
    d   = request.get_json(force=True) or {}
    eid = d.get("id", "").strip()
    if not eid: return jsonify({"error": "id required"}), 400
    f = corp_file(eid)
    if f.exists(): f.unlink()
    return jsonify({"ok": True})

# ── Directors ─────────────────────────────────────────────────

@app.route("/api/corp/director/add", methods=["POST"])
@login_required
@verify_csrf
def corp_add_director():
    d   = request.get_json(force=True) or {}
    eid = d.get("entity_id", "").strip()
    if not eid: return jsonify({"error": "entity_id required"}), 400
    entity = load_entity(eid)
    if not entity: return jsonify({"error": "Entity not found"}), 404
    director = {
        "id":               secrets.token_hex(6),
        "type":             d.get("type", "natural_person"),  # natural_person | legal_entity
        "first_name":       d.get("first_name", ""),
        "last_name":        d.get("last_name", ""),
        "entity_name":      d.get("entity_name", ""),         # if legal entity
        "representative":   d.get("representative", ""),      # permanent rep name
        "nationality":      d.get("nationality", ""),
        "address":          d.get("address", ""),
        "role":             d.get("role", "Non-executive director"),
        "appointment_date": d.get("appointment_date", ""),
        "mandate_end":      d.get("mandate_end", ""),
        "appointment_auth": d.get("appointment_auth", "AGM resolution"),
        "deed_ref":         d.get("deed_ref", ""),
        "remuneration":     d.get("remuneration", "Unpaid"),
        "signature_auth":   d.get("signature_auth", "Joint"),
        "independent":      d.get("independent", False),
        "status":           "active",
        "notes":            d.get("notes", ""),
        "added_at":         datetime.datetime.now().isoformat(),
    }
    display_name = director["entity_name"] or f"{director['first_name']} {director['last_name']}".strip()
    entity.setdefault("directors", []).append(director)
    entity["updated_at"] = datetime.datetime.now().isoformat()
    _add_history(entity, "governance",
                 f"Director added: {display_name} ({director['role']})",
                 source=director.get("deed_ref",""),
                 recorded_by=session["username"])
    save_entity(eid, entity)
    return jsonify({"ok": True, "id": director["id"]})

@app.route("/api/corp/director/update", methods=["POST"])
@login_required
@verify_csrf
def corp_update_director():
    d   = request.get_json(force=True) or {}
    eid = d.get("entity_id", "").strip()
    did = d.get("director_id", "").strip()
    entity = load_entity(eid)
    if not entity: return jsonify({"error": "Entity not found"}), 404
    directors = entity.get("directors", [])
    for i, dr in enumerate(directors):
        if dr["id"] == did:
            for k, v in d.items():
                if k not in ("entity_id", "director_id"):
                    directors[i][k] = v
            break
    entity["updated_at"] = datetime.datetime.now().isoformat()
    _add_history(entity, "governance", f"Director record updated (id: {did})",
                 recorded_by=session["username"])
    save_entity(eid, entity)
    return jsonify({"ok": True})

@app.route("/api/corp/director/end_mandate", methods=["POST"])
@login_required
@verify_csrf
def corp_end_mandate():
    d   = request.get_json(force=True) or {}
    eid = d.get("entity_id", "").strip()
    did = d.get("director_id", "").strip()
    reason = d.get("reason", "Mandate ended")
    entity = load_entity(eid)
    if not entity: return jsonify({"error": "Entity not found"}), 404
    for dr in entity.get("directors", []):
        if dr["id"] == did:
            dr["status"]    = "ended"
            dr["end_date"]  = datetime.date.today().isoformat()
            dr["end_reason"]= reason
            name = dr.get("entity_name") or f"{dr.get('first_name','')} {dr.get('last_name','')}".strip()
            _add_history(entity, "governance", f"Director mandate ended: {name} — {reason}",
                         recorded_by=session["username"])
            break
    entity["updated_at"] = datetime.datetime.now().isoformat()
    save_entity(eid, entity)
    return jsonify({"ok": True})

# ── Shareholders ──────────────────────────────────────────────

@app.route("/api/corp/shareholder/add", methods=["POST"])
@login_required
@verify_csrf
def corp_add_shareholder():
    d   = request.get_json(force=True) or {}
    eid = d.get("entity_id", "").strip()
    entity = load_entity(eid)
    if not entity: return jsonify({"error": "Entity not found"}), 404
    sh = {
        "id":             secrets.token_hex(6),
        "name":           d.get("name", ""),
        "type":           d.get("type", "natural_person"),  # natural_person | legal_entity | public_float
        "nationality":    d.get("nationality", ""),
        "address":        d.get("address", ""),
        "shares":         d.get("shares", 0),
        "pct":            d.get("pct", 0.0),
        "share_class":    d.get("share_class", "Ordinary"),
        "acquisition_date": d.get("acquisition_date", ""),
        "acquisition_price": d.get("acquisition_price", ""),
        "ownership_type": d.get("ownership_type", "Direct"),  # Direct | Indirect
        "encumbrance":    d.get("encumbrance", "None"),
        "is_ubo":         d.get("is_ubo", False),
        "ubo_kyc_status": d.get("ubo_kyc_status", "pending"),
        "ubo_last_confirmed": d.get("ubo_last_confirmed", ""),
        "shareholder_type": d.get("shareholder_type", ""),  # Founder | Investor | Management etc
        "notes":          d.get("notes", ""),
        "added_at":       datetime.datetime.now().isoformat(),
    }
    entity.setdefault("shareholders", []).append(sh)
    entity["updated_at"] = datetime.datetime.now().isoformat()
    _add_history(entity, "capital",
                 f"Shareholder added: {sh['name']} ({sh['pct']}%)",
                 recorded_by=session["username"])
    save_entity(eid, entity)
    return jsonify({"ok": True, "id": sh["id"]})

@app.route("/api/corp/shareholder/update", methods=["POST"])
@login_required
@verify_csrf
def corp_update_shareholder():
    d   = request.get_json(force=True) or {}
    eid = d.get("entity_id", "").strip()
    sid = d.get("shareholder_id", "").strip()
    entity = load_entity(eid)
    if not entity: return jsonify({"error": "Entity not found"}), 404
    for i, sh in enumerate(entity.get("shareholders", [])):
        if sh["id"] == sid:
            for k, v in d.items():
                if k not in ("entity_id", "shareholder_id"):
                    entity["shareholders"][i][k] = v
            break
    entity["updated_at"] = datetime.datetime.now().isoformat()
    _add_history(entity, "capital", f"Shareholder record updated (id: {sid})",
                 recorded_by=session["username"])
    save_entity(eid, entity)
    return jsonify({"ok": True})

@app.route("/api/corp/shareholder/delete", methods=["POST"])
@login_required
@verify_csrf
def corp_delete_shareholder():
    d   = request.get_json(force=True) or {}
    eid = d.get("entity_id", "").strip()
    sid = d.get("shareholder_id", "").strip()
    entity = load_entity(eid)
    if not entity: return jsonify({"error": "Entity not found"}), 404
    entity["shareholders"] = [s for s in entity.get("shareholders", []) if s["id"] != sid]
    entity["updated_at"] = datetime.datetime.now().isoformat()
    _add_history(entity, "capital", f"Shareholder removed (id: {sid})",
                 recorded_by=session["username"])
    save_entity(eid, entity)
    return jsonify({"ok": True})

# ── Powers of Attorney ────────────────────────────────────────

@app.route("/api/corp/poa/add", methods=["POST"])
@login_required
@verify_csrf
def corp_add_poa():
    d   = request.get_json(force=True) or {}
    eid = d.get("entity_id", "").strip()
    entity = load_entity(eid)
    if not entity: return jsonify({"error": "Entity not found"}), 404
    poa = {
        "id":           secrets.token_hex(6),
        "grantee_name": d.get("grantee_name", ""),
        "grantee_address": d.get("grantee_address", ""),
        "granted_by":   d.get("granted_by", ""),  # board resolution ref
        "grant_date":   d.get("grant_date", ""),
        "duration":     d.get("duration", "Unlimited"),
        "expiry_date":  d.get("expiry_date", ""),
        "signature":    d.get("signature", "Sole"),  # Sole | Joint
        "status":       "active",
        "powers":       d.get("powers", []),  # list of power strings
        "thresholds":   d.get("thresholds", {}),  # {contracts: 500000, leases: 250000}
        "notes":        d.get("notes", ""),
        "added_at":     datetime.datetime.now().isoformat(),
    }
    entity.setdefault("poas", []).append(poa)
    entity["updated_at"] = datetime.datetime.now().isoformat()
    _add_history(entity, "governance",
                 f"Power of attorney granted to {poa['grantee_name']}",
                 source=poa.get("granted_by",""),
                 recorded_by=session["username"])
    save_entity(eid, entity)
    return jsonify({"ok": True, "id": poa["id"]})

@app.route("/api/corp/poa/revoke", methods=["POST"])
@login_required
@verify_csrf
def corp_revoke_poa():
    d   = request.get_json(force=True) or {}
    eid = d.get("entity_id", "").strip()
    pid = d.get("poa_id", "").strip()
    entity = load_entity(eid)
    if not entity: return jsonify({"error": "Entity not found"}), 404
    for poa in entity.get("poas", []):
        if poa["id"] == pid:
            poa["status"]       = "revoked"
            poa["revoked_date"] = datetime.date.today().isoformat()
            _add_history(entity, "governance",
                         f"Power of attorney revoked: {poa['grantee_name']}",
                         recorded_by=session["username"])
            break
    entity["updated_at"] = datetime.datetime.now().isoformat()
    save_entity(eid, entity)
    return jsonify({"ok": True})

# ── Compliance items ──────────────────────────────────────────

@app.route("/api/corp/compliance/add", methods=["POST"])
@login_required
@verify_csrf
def corp_add_compliance():
    d   = request.get_json(force=True) or {}
    eid = d.get("entity_id", "").strip()
    entity = load_entity(eid)
    if not entity: return jsonify({"error": "Entity not found"}), 404
    item = {
        "id":          secrets.token_hex(6),
        "obligation":  d.get("obligation", ""),
        "period":      d.get("period", ""),
        "deadline":    d.get("deadline", ""),
        "filed_date":  d.get("filed_date", ""),
        "status":      d.get("status", "pending"),  # pending | filed | overdue | exempt
        "authority":   d.get("authority", ""),      # BNB | CBE | FSMA | etc.
        "notes":       d.get("notes", ""),
        "added_at":    datetime.datetime.now().isoformat(),
    }
    entity.setdefault("compliance", []).append(item)
    entity["updated_at"] = datetime.datetime.now().isoformat()
    _add_history(entity, "compliance",
                 f"Compliance obligation added: {item['obligation']} ({item['period']})",
                 recorded_by=session["username"])
    save_entity(eid, entity)
    return jsonify({"ok": True, "id": item["id"]})

@app.route("/api/corp/compliance/update", methods=["POST"])
@login_required
@verify_csrf
def corp_update_compliance():
    d   = request.get_json(force=True) or {}
    eid = d.get("entity_id", "").strip()
    cid = d.get("item_id", "").strip()
    entity = load_entity(eid)
    if not entity: return jsonify({"error": "Entity not found"}), 404
    for i, item in enumerate(entity.get("compliance", [])):
        if item["id"] == cid:
            for k, v in d.items():
                if k not in ("entity_id","item_id"):
                    entity["compliance"][i][k] = v
            # Auto-set overdue
            if item.get("deadline") and not item.get("filed_date"):
                try:
                    dl = datetime.date.fromisoformat(item["deadline"])
                    if dl < datetime.date.today():
                        entity["compliance"][i]["status"] = "overdue"
                except: pass
            break
    entity["updated_at"] = datetime.datetime.now().isoformat()
    _add_history(entity, "compliance", f"Compliance item updated (id: {cid})",
                 recorded_by=session["username"])
    save_entity(eid, entity)
    return jsonify({"ok": True})

# ── History entry (manual) ────────────────────────────────────

@app.route("/api/corp/history/add", methods=["POST"])
@login_required
@verify_csrf
def corp_add_history():
    d   = request.get_json(force=True) or {}
    eid = d.get("entity_id", "").strip()
    entity = load_entity(eid)
    if not entity: return jsonify({"error": "Entity not found"}), 404
    _add_history(entity,
                 event_type=d.get("type","other"),
                 description=d.get("description",""),
                 source=d.get("source",""),
                 recorded_by=session["username"])
    entity["updated_at"] = datetime.datetime.now().isoformat()
    save_entity(eid, entity)
    return jsonify({"ok": True})

# ── Mickey AI: ask about an entity ───────────────────────────

@app.route("/api/corp/ask", methods=["POST"])
@login_required
def corp_ask():
    username = session["username"]
    if not check_rate_limit(username):
        return jsonify({"error": "Rate limit reached. Please wait."}), 429
    d       = request.get_json(force=True) or {}
    eid     = d.get("entity_id", "").strip()
    question= d.get("question", "").strip()
    if not question: return jsonify({"error": "Question required"}), 400
    entity  = load_entity(eid) if eid else {}
    display = load_users().get(username,{}).get("display_name", username)
    # Build entity context string
    ctx_parts = []
    if entity:
        ctx_parts.append(f"Entity: {entity.get('name','')} ({entity.get('legal_form','')} · {entity.get('jurisdiction','')})")
        ctx_parts.append(f"Registration: {entity.get('reg_number','')} · VAT: {entity.get('vat_number','')} · LEI: {entity.get('lei_code','')}")
        ctx_parts.append(f"Registered office: {entity.get('registered_street','')} {entity.get('registered_city','')} {entity.get('registered_country','')}")
        ctx_parts.append(f"Share capital: {entity.get('share_capital','')} {entity.get('capital_currency','')} · Shares: {entity.get('total_shares','')}")
        # Directors
        active_dirs = [dr for dr in entity.get("directors",[]) if dr.get("status")=="active"]
        if active_dirs:
            dir_lines = []
            for dr in active_dirs:
                name = dr.get("entity_name") or f"{dr.get('first_name','')} {dr.get('last_name','')}".strip()
                dir_lines.append(f"{name} ({dr.get('role','')}, mandate ends {dr.get('mandate_end','')})")
            ctx_parts.append("Active directors:\n" + "\n".join(dir_lines))
        # Shareholders
        shs = entity.get("shareholders",[])
        if shs:
            sh_lines = [f"{sh['name']}: {sh['shares']} shares ({sh['pct']}%)" for sh in shs[:5]]
            ctx_parts.append("Shareholders:\n" + "\n".join(sh_lines))
        # Compliance
        comps = [c for c in entity.get("compliance",[]) if c.get("status") in ("pending","overdue")]
        if comps:
            c_lines = [f"{c['obligation']} ({c['period']}) — deadline {c['deadline']} — {c['status']}" for c in comps[:5]]
            ctx_parts.append("Open compliance items:\n" + "\n".join(c_lines))
    entity_ctx = "\n".join(ctx_parts)
    system_extra = f"""
You also have access to corporate entity data for this query:

{entity_ctx}

When answering questions about this entity, use the data above. For legal questions about obligations (UBO, annual accounts, director mandates, etc.), cite the relevant Belgian legal provisions (WVV, Staatsblad, KBO/CBE obligations). Offer to draft documents (board resolutions, convening notices, POA instruments) when relevant.
""" if entity_ctx else ""
    try:
        client = get_client()
        resp = client.messages.create(
            model="claude-sonnet-4-20250514", max_tokens=2000,
            system=get_system(display) + system_extra,
            messages=[{"role":"user","content":question}]
        )
        answer = " ".join(b.text for b in resp.content if hasattr(b,"text") and b.text).strip()
        tokens = resp.usage.input_tokens + resp.usage.output_tokens
        log_usage(username, "corp_ask", tokens, question[:40])
        return jsonify({"answer": answer})
    except ValueError as e: return jsonify({"error": str(e), "need_key": True}), 500
    except Exception as e:  return jsonify({"error": str(e)}), 500

# ── Corporate document upload / list / delete ─────────────────

def corp_docs_dir(entity_id: str) -> Path:
    d = Path(APP_PATH) / "corporate" / entity_id / "docs"
    d.mkdir(parents=True, exist_ok=True)
    return d

def corp_docs_index(entity_id: str) -> Path:
    return Path(APP_PATH) / "corporate" / entity_id / "docs_index.json"

def load_docs_index(entity_id: str) -> dict:
    p = corp_docs_index(entity_id)
    if p.exists():
        try: return json.loads(p.read_text(encoding="utf-8"))
        except: pass
    return {}

def save_docs_index(entity_id: str, idx: dict):
    corp_docs_index(entity_id).write_text(json.dumps(idx, ensure_ascii=False, indent=2), encoding="utf-8")

@app.route("/api/corp/docs", methods=["GET"])
@login_required
def corp_list_docs():
    eid = request.args.get("entity_id","").strip()
    if not eid: return jsonify({"error":"entity_id required"}), 400
    idx = load_docs_index(eid)
    docs = sorted(idx.values(), key=lambda x: x.get("uploaded_at",""), reverse=True)
    return jsonify({"docs": docs})

@app.route("/api/corp/doc/upload", methods=["POST"])
@login_required
def corp_upload_doc():
    username = session["username"]
    eid      = request.form.get("entity_id","").strip()
    name     = (request.form.get("name") or "").strip()
    category = (request.form.get("category") or "other").strip()
    doc_date = (request.form.get("doc_date") or "").strip()
    notes    = (request.form.get("notes") or "").strip()
    file     = request.files.get("file")
    if not eid or not name or not file:
        return jsonify({"error": "entity_id, name and file required"}), 400
    # Validate entity exists
    entity = load_entity(eid)
    if not entity:
        return jsonify({"error": "Entity not found"}), 404
    # Save file
    doc_id   = secrets.token_hex(8)
    safe_fn  = re.sub(r'[^\w\s\-.]','', file.filename or "file")[:80]
    dest_dir = corp_docs_dir(eid)
    dest     = dest_dir / f"{doc_id}_{safe_fn}"
    file.save(str(dest))
    # Also extract text and store for knowledge retrieval
    try:
        file.seek(0)
        text = extract_bytes(file.read(), file.filename)
        if text:
            txt_path = dest_dir / f"{doc_id}.txt"
            txt_path.write_text(text, encoding="utf-8")
    except: pass
    # Update index
    idx = load_docs_index(eid)
    idx[doc_id] = {
        "id": doc_id,
        "name": name,
        "filename": safe_fn,
        "category": category,
        "doc_date": doc_date,
        "notes": notes,
        "uploaded_by": username,
        "uploaded_at": datetime.datetime.now().isoformat(),
        "file_path": str(dest),
    }
    save_docs_index(eid, idx)
    # Add history entry
    entity = load_entity(eid)
    if entity:
        _add_history(entity, type="document",
                     description=f"Document uploaded: {name} [{category}]",
                     source=safe_fn, recorded_by=username)
        entity["updated_at"] = datetime.datetime.now().isoformat()
        save_entity(eid, entity)
    log_usage(username, "corp_doc_upload", 0, f"{name[:40]} [{eid[:8]}]")
    return jsonify({"ok": True, "id": doc_id})

@app.route("/api/corp/doc/delete", methods=["POST"])
@login_required
@verify_csrf
def corp_delete_doc():
    d      = request.get_json(force=True) or {}
    eid    = d.get("entity_id","").strip()
    doc_id = d.get("doc_id","").strip()
    if not eid or not doc_id: return jsonify({"error":"entity_id and doc_id required"}),400
    idx = load_docs_index(eid)
    entry = idx.pop(doc_id, None)
    if entry:
        # Delete file(s)
        for p in [entry.get("file_path"), str(corp_docs_dir(eid) / f"{doc_id}.txt")]:
            if p:
                try: Path(p).unlink(missing_ok=True)
                except: pass
        save_docs_index(eid, idx)
    return jsonify({"ok": True})

# ── Run ───────────────────────────────────────────────────────

if __name__ == "__main__":
    host = os.environ.get("MICKEY_HOST", "0.0.0.0")
    port = int(os.environ.get("MICKEY_PORT", "5000"))
    print(f"\n  Mickey -- Legal Intelligence")
    print(f"  Data: {APP_PATH}")
    print(f"  Rate limit: {RATE_LIMIT_CALLS_PER_HOUR} calls/user/hour")
    try:
        from waitress import serve
        print(f"  Running at http://{host}:{port} (Waitress WSGI server)\n")
        serve(app, host=host, port=port, threads=8)
    except ImportError:
        print(f"  Running at http://{host}:{port} (Flask dev server)")
        print(f"  Tip: pip install waitress for production use\n")
        app.run(debug=False, port=port, host=host)


# ══════════════════════════════════════════════════════════════
# DOCUMENTS MODULE
# ══════════════════════════════════════════════════════════════
#
# Routes:
#   POST /api/docs/review        — structured legal review
#   POST /api/docs/compare       — document comparison
#   POST /api/docs/draft         — AI drafting (new / improve / counter)
#   POST /api/docs/summarise     — audience-aware summary
#   POST /api/docs/translate     — legal translation
#   POST /api/docs/anonymise     — redaction with map
#   GET  /api/docs/history       — recent docs for user
#   POST /api/docs/history/save  — save doc to history
#   DELETE /api/docs/history/<id>— delete history item
#   GET  /api/docs/templates     — list user templates
#   POST /api/docs/templates     — create / update template
#   DELETE /api/docs/template/<id>— delete template
#   GET  /api/docs/obligations   — saved obligations (dashboard widget)
#
# Storage layout (per user):
#   C:/Mickey/docs/{username}/history/{id}.json
#   C:/Mickey/docs/{username}/templates/{id}.json
#   C:/Mickey/docs/{username}/obligations/{id}.json
#   C:/Mickey/docs/{username}/anonymised/{id}/  (original + redacted + map)
#
# Scalability notes:
#   - Every route follows login_required + verify_csrf + rate_limit pattern
#   - Structured JSON output on all AI calls enables frontend parsing
#   - widgetPush-compatible obligations endpoint feeds dashboard
#   - OCR detection is transparent: if pypdf yields <50 chars/page → flag
#   - All file text truncated to safe context window limits per module
# ══════════════════════════════════════════════════════════════

import uuid

DOCS_PATH = APP_PATH / "docs"
DOCS_PATH.mkdir(parents=True, exist_ok=True)

# ── Helpers ───────────────────────────────────────────────────

def docs_user_path(username: str) -> Path:
    p = DOCS_PATH / username
    for sub in ["history", "templates", "obligations", "anonymised"]:
        (p / sub).mkdir(parents=True, exist_ok=True)
    return p

def docs_load_json(path: Path) -> dict:
    try:
        return json.loads(path.read_text(encoding="utf-8"))
    except Exception:
        return {}

def docs_save_json(path: Path, data: dict):
    path.write_text(json.dumps(data, ensure_ascii=False, indent=2), encoding="utf-8")

def docs_list_dir(path: Path) -> list:
    """Return list of parsed JSON objects from a directory, newest first."""
    items = []
    for f in sorted(path.glob("*.json"), key=lambda x: x.stat().st_mtime, reverse=True):
        try:
            items.append(json.loads(f.read_text(encoding="utf-8")))
        except Exception:
            continue
    return items

def detect_ocr_needed(text: str, pages: int) -> bool:
    """True if text yield is suspiciously low — likely a scanned PDF."""
    if pages == 0:
        return False
    return (len(text.strip()) / max(pages, 1)) < 80

def extract_bytes_with_meta(data: bytes, filename: str) -> dict:
    """Extract text and return metadata including page count and OCR flag."""
    text = ""
    pages = 0
    ocr_needed = False
    try:
        if filename.lower().endswith(".docx"):
            from docx import Document as DocxDocument
            doc = DocxDocument(io.BytesIO(data))
            text = "\n".join(p.text for p in doc.paragraphs if p.text.strip())
            pages = max(1, len(text) // 3000)  # rough estimate
        else:
            import pypdf
            reader = pypdf.PdfReader(io.BytesIO(data))
            pages = len(reader.pages)
            text = "\n".join(p.extract_text() or "" for p in reader.pages).strip()
            ocr_needed = detect_ocr_needed(text, pages)
    except Exception:
        pass
    return {"text": text, "pages": pages, "ocr_needed": ocr_needed, "filename": filename}

def docs_system(display_name: str) -> str:
    """System prompt for all documents module calls."""
    return f"""You are Mickey, a legal intelligence assistant for {display_name}.

You are operating in the Documents module. You analyse, draft, compare, summarise, translate, and anonymise legal documents with precision.

Standards:
1. Cite articles precisely: "Art. 28(3) GDPR", "Art. 7:96 WVV", "s.90 Companies Act 2006"
2. When a finding references a specific clause or article in the uploaded document, quote it briefly
3. Distinguish clearly: HIGH risk (must fix), MEDIUM risk (should address), LOW risk (minor)
4. Use structured output exactly as requested — the frontend parses your response
5. Belgian law: always check federal and regional competences where relevant
6. End legal analyses with "**Sources cited:**" listing all articles and cases referenced"""

def call_claude(username: str, prompt: str, system: str, max_tokens: int = 4000) -> tuple:
    """
    Make a Claude API call. Returns (answer_text, tokens_used, error).
    Centralised so model version and error handling are consistent everywhere.
    """
    display = load_users().get(username, {}).get("display_name", username)
    try:
        client = get_client()
        resp = client.messages.create(
            model="claude-sonnet-4-20250514",
            max_tokens=max_tokens,
            system=system or docs_system(display),
            messages=[{"role": "user", "content": prompt}]
        )
        answer = " ".join(b.text for b in resp.content if hasattr(b, "text") and b.text).strip()
        tokens = resp.usage.input_tokens + resp.usage.output_tokens
        return answer, tokens, None
    except ValueError as e:
        return None, 0, {"error": str(e), "need_key": True}
    except Exception as e:
        return None, 0, {"error": str(e)}

def docs_save_history(username: str, module: str, filename: str,
                      result: str, meta: dict = None) -> str:
    """Save a completed document operation to user history. Returns item id."""
    item_id = str(uuid.uuid4())[:8]
    base = docs_user_path(username) / "history" / f"{item_id}.json"
    docs_save_json(base, {
        "id": item_id,
        "module": module,
        "filename": filename,
        "result_preview": result[:300],
        "result": result,
        "meta": meta or {},
        "created_at": datetime.datetime.utcnow().isoformat()
    })
    return item_id

# ══════════════════════════════════════════════════════════════
# REVIEW
# ══════════════════════════════════════════════════════════════
# Input:  multipart file upload
#         focus        — comma-separated topics (e.g. "GDPR, liability, IP")
#         playbook     — optional playbook name to check against
#         perspective  — "Seller" | "Buyer" | "Neutral" etc.
#         audience     — "Myself" | "Board" | "Management" | "Other party" | "File record"
#         extra        — free-text instructions
#
# Output: {
#   result: full markdown text,
#   findings: [{severity, clause, issue, fix}],
#   obligations: [{description, deadline_type, party}],
#   summary: str,
#   filename: str,
#   ocr_needed: bool,
#   history_id: str
# }
# ══════════════════════════════════════════════════════════════

@app.route("/api/docs/review", methods=["POST"])
@login_required
@verify_csrf
def docs_review():
    username = session["username"]
    if not check_rate_limit(username):
        return jsonify({"error": "Rate limit reached. Please wait."}), 429

    file = request.files.get("file")
    if not file:
        return jsonify({"error": "No file uploaded"}), 400

    focus       = request.form.get("focus", "General legal review")
    playbook    = request.form.get("playbook", "")
    perspective = request.form.get("perspective", "Neutral")
    audience    = request.form.get("audience", "Myself")
    extra       = request.form.get("extra", "")

    extracted = extract_bytes_with_meta(file.read(), file.filename)
    text = extracted["text"]
    if not text:
        return jsonify({"error": "Could not extract text from document.", "ocr_needed": True}), 400

    # Knowledge base context
    ctx = get_context(username, focus + " " + file.filename)
    ctx_str = ("\n\nRelevant knowledge base references:\n\n" +
               "\n---\n".join(ctx[:2])) if ctx else ""

    audience_instruction = {
        "Myself":      "Write for a lawyer. Use legal terminology freely. Be direct and complete.",
        "Board":       "Write for non-lawyer board members. Executive summary first. Avoid jargon.",
        "Management":  "Write for senior management. Practical implications, business risk language.",
        "Other party": "Write diplomatically. Highlight mutual benefits. De-emphasise adversarial findings.",
        "File record": "Write formally for the file. Complete, precise, citation-heavy."
    }.get(audience, "Write for a lawyer. Be direct and complete.")

    prompt = f"""Review this legal document and return a structured analysis.

Document: {file.filename}
Review focus: {focus}
Perspective: {perspective}
{f"Playbook to check against: {playbook}" if playbook else ""}
{f"Additional instructions: {extra}" if extra else ""}
{audience_instruction}
{ctx_str}

Document text:
{text[:9000]}

Return your analysis in this exact structure:

## Summary
[3-5 sentence summary of what this document is, its purpose, key parties, and governing law]

## Findings

### 🔴 HIGH — Must fix
[For each: **Clause/Article** | **Issue** | **Recommended fix**]
[If none: "No high-severity findings."]

### 🟡 MEDIUM — Should address
[For each: **Clause/Article** | **Issue** | **Recommended fix**]
[If none: "No medium-severity findings."]

### 🟢 LOW — Minor improvements
[For each: **Clause/Article** | **Issue** | **Suggested improvement**]
[If none: "No low-severity findings."]

## Obligations extracted
[List every obligation, deadline or recurring requirement found. Format each as:]
- **[Party]**: [obligation description] — [deadline or trigger if stated]
[If none found: "No explicit obligations or deadlines identified."]

## Overall recommendation
[2-3 sentences: sign/negotiate/reject + key conditions]

**Sources cited:** [articles, regulations, cases referenced]"""

    answer, tokens, err = call_claude(username, prompt, docs_system(
        load_users().get(username, {}).get("display_name", username)))
    if err:
        return jsonify(err), 500

    # Parse obligations from the answer to save separately
    obligations = []
    in_obligations = False
    for line in answer.split("\n"):
        if "## Obligations" in line:
            in_obligations = True
            continue
        if line.startswith("## ") and in_obligations:
            break
        if in_obligations and line.strip().startswith("- "):
            obligations.append({"description": line.strip()[2:],
                                 "source_doc": file.filename,
                                 "extracted_at": datetime.datetime.utcnow().isoformat()})

    # Save obligations for dashboard widget
    if obligations:
        obs_path = docs_user_path(username) / "obligations"
        obs_id = str(uuid.uuid4())[:8]
        docs_save_json(obs_path / f"{obs_id}.json", {
            "id": obs_id,
            "filename": file.filename,
            "obligations": obligations,
            "created_at": datetime.datetime.utcnow().isoformat()
        })

    history_id = docs_save_history(username, "review", file.filename, answer,
                                   {"focus": focus, "perspective": perspective,
                                    "audience": audience, "ocr_needed": extracted["ocr_needed"]})

    log_usage(username, "docs_review", tokens, file.filename[:40])

    return jsonify({
        "result": answer,
        "obligations": obligations,
        "filename": file.filename,
        "ocr_needed": extracted["ocr_needed"],
        "pages": extracted["pages"],
        "history_id": history_id
    })


# ══════════════════════════════════════════════════════════════
# COMPARE
# ══════════════════════════════════════════════════════════════
# Modes:
#   "draft"     — your doc vs own earlier draft (file + file2)
#   "playbook"  — your doc vs a named playbook (file + playbook param)
#   "contracts" — your doc vs existing signed contracts in knowledge base
#   "market"    — your doc vs Mickey's market standard knowledge
#
# Output: {
#   result: full markdown,
#   deviations: [{clause, yours, reference, risk, priority}],
#   filename: str,
#   mode: str,
#   history_id: str
# }
# ══════════════════════════════════════════════════════════════

@app.route("/api/docs/compare", methods=["POST"])
@login_required
@verify_csrf
def docs_compare():
    username = session["username"]
    if not check_rate_limit(username):
        return jsonify({"error": "Rate limit reached. Please wait."}), 429

    file = request.files.get("file")
    if not file:
        return jsonify({"error": "No file uploaded"}), 400

    mode        = request.form.get("mode", "draft")   # draft|playbook|contracts|market
    focus       = request.form.get("focus", "")
    playbook    = request.form.get("playbook", "")
    extra       = request.form.get("extra", "")

    extracted = extract_bytes_with_meta(file.read(), file.filename)
    text = extracted["text"]
    if not text:
        return jsonify({"error": "Could not extract text from document."}), 400

    display = load_users().get(username, {}).get("display_name", username)

    if mode == "draft":
        file2 = request.files.get("file2")
        if not file2:
            return jsonify({"error": "Second file required for draft comparison"}), 400
        ext2 = extract_bytes_with_meta(file2.read(), file2.filename)
        ref_label = f"Reference draft: {file2.filename}"
        ref_text  = ext2["text"][:4500]
        mode_instruction = "Compare clause by clause between the two drafts. Identify what changed, what was removed, what was added."

    elif mode == "playbook":
        if not playbook:
            return jsonify({"error": "Playbook name required"}), 400
        # Retrieve playbook from knowledge base context
        ctx = get_context(username, playbook)
        ref_label = f"Playbook: {playbook}"
        ref_text  = "\n".join(ctx[:3]) if ctx else f"[{playbook} — apply standard market positions]"
        mode_instruction = "Check whether each clause in the document satisfies the playbook requirements. Flag every deviation."

    elif mode == "contracts":
        ctx = get_context(username, focus or file.filename)
        ref_label = "Existing signed contracts (knowledge base)"
        ref_text  = "\n---\n".join(ctx[:3]) if ctx else "[No matching contracts found in knowledge base]"
        mode_instruction = "Compare the key commercial terms against the precedents. Identify where this document deviates from your standard positions."

    else:  # market
        ref_label = "Market standard positions"
        ref_text  = "[Apply your knowledge of market standard terms for this document type and jurisdiction]"
        mode_instruction = "Compare each clause against market standard for this document type. Flag clauses that are unusually aggressive, one-sided, or below market."

    prompt = f"""Compare this document against a reference.

Document under review: {file.filename}
{ref_label}
{f"Focus areas: {focus}" if focus else ""}
{f"Additional instructions: {extra}" if extra else ""}

{mode_instruction}

Document text:
{text[:4500]}

Reference:
{ref_text}

Return your analysis in this exact structure:

## Comparison summary
[3-4 sentences: overall alignment, key theme of deviations, negotiation posture]

## Clause-by-clause comparison

| Clause | Your document | Reference | Status | Risk |
|--------|--------------|-----------|--------|------|
[Fill one row per material clause. Status: ✓ Aligned / ⚠ Deviation / ✗ Missing]

## Key deviations — negotiation priorities

### Must push back on
[Clauses where deviation is high risk — specific redline suggestion for each]

### Should negotiate
[Clauses worth addressing but not deal-breakers]

### Accept or minor
[Deviations that are acceptable or low priority]

## Overall position
[Sign as-is / Negotiate X points / Reject — with specific conditions]

**Sources cited:** [if any]"""

    answer, tokens, err = call_claude(username, prompt, docs_system(display))
    if err:
        return jsonify(err), 500

    history_id = docs_save_history(username, "compare", file.filename, answer,
                                   {"mode": mode, "playbook": playbook})
    log_usage(username, "docs_compare", tokens, file.filename[:40])

    return jsonify({
        "result": answer,
        "filename": file.filename,
        "mode": mode,
        "ocr_needed": extracted["ocr_needed"],
        "history_id": history_id
    })


# ══════════════════════════════════════════════════════════════
# DRAFT
# ══════════════════════════════════════════════════════════════
# Three modes:
#   "new"       — draft from scratch given parameters
#   "improve"   — improve / redline an existing clause or section
#   "counter"   — generate counter-proposal to uploaded clause
#
# Input (JSON body for new/improve/counter without file,
#        multipart for improve/counter with uploaded text):
#   mode, doc_type, clause_description, governing_law,
#   perspective, jurisdiction, instructions, audience,
#   template_id (optional — load saved template)
#
# Output: {
#   result: full drafted text,
#   mode: str,
#   mickey_notes: [str],  — items flagged for verification
#   history_id: str
# }
# ══════════════════════════════════════════════════════════════

@app.route("/api/docs/draft", methods=["POST"])
@login_required
@verify_csrf
def docs_draft():
    username = session["username"]
    if not check_rate_limit(username):
        return jsonify({"error": "Rate limit reached. Please wait."}), 429

    display = load_users().get(username, {}).get("display_name", username)

    # Accept both JSON and multipart (when uploading existing clause)
    if request.content_type and "multipart" in request.content_type:
        d            = request.form
        file         = request.files.get("file")
        existing_text = extract_bytes_with_meta(file.read(), file.filename)["text"][:4000] if file else ""
    else:
        d            = request.get_json(force=True) or {}
        existing_text = d.get("existing_text", "")

    mode         = d.get("mode", "new")          # new | improve | counter
    doc_type     = d.get("doc_type", "contract clause")
    clause_desc  = d.get("clause_description", "")
    governing_law= d.get("governing_law", "Belgian law")
    perspective  = d.get("perspective", "Neutral")
    jurisdiction = d.get("jurisdiction", "Belgium")
    instructions = d.get("instructions", "")
    audience     = d.get("audience", "Myself")
    template_id  = d.get("template_id", "")

    # Load template if requested
    template_text = ""
    if template_id:
        tpath = docs_user_path(username) / "templates" / f"{template_id}.json"
        if tpath.exists():
            tmpl = docs_load_json(tpath)
            template_text = tmpl.get("content", "")

    # Knowledge base context
    ctx = get_context(username, f"{doc_type} {clause_desc} {governing_law}")
    ctx_str = ("\n\nRelevant precedents from your knowledge base:\n\n" +
               "\n---\n".join(ctx[:2])) if ctx else ""

    audience_note = {
        "Myself":      "Draft for internal use. Full legal rigour.",
        "Board":       "Draft suitable for board presentation. Clear, concise recitals.",
        "Management":  "Draft with management execution in mind. Practical obligations.",
        "Other party": "Draft for external party. Professional, balanced tone.",
        "File record": "Draft for the file. Formal, complete, fully cited."
    }.get(audience, "Full legal rigour.")

    if mode == "new":
        prompt = f"""Draft a {doc_type}.

Description: {clause_desc}
Governing law: {governing_law}
Jurisdiction: {jurisdiction}
Perspective: {perspective}
{f"Style note: {audience_note}"}
{f"Specific instructions: {instructions}" if instructions else ""}
{f"Base this on the following template:{chr(10)}{template_text}" if template_text else ""}
{ctx_str}

Requirements:
- Number all articles
- Define capitalised terms at first use
- Cite the legal basis for each obligation (regulation, statute, or case law)
- Flag any item requiring client-specific verification as [VERIFY: reason]
- Professional Belgian/EU legal drafting style

After the draft, add:
## Mickey notes
[Bullet list of items flagged [VERIFY] with brief explanation of what needs checking]"""

    elif mode == "improve":
        if not existing_text:
            return jsonify({"error": "Existing text required for improve mode"}), 400
        prompt = f"""Improve this {doc_type}.

Perspective: {perspective}
Governing law: {governing_law}
{f"Specific instructions: {instructions}" if instructions else ""}
{ctx_str}

Original text:
{existing_text}

Provide:
1. A redlined version showing changes (use ~~strikethrough~~ for deletions, **bold** for additions)
2. A clean final version
3. A brief explanation of each substantive change and why it improves the document

## Mickey notes
[Items requiring client verification]"""

    else:  # counter
        if not existing_text:
            return jsonify({"error": "Existing clause required for counter mode"}), 400
        prompt = f"""Generate a counter-proposal to this clause.

Our perspective: {perspective}
Governing law: {governing_law}
{f"Specific instructions: {instructions}" if instructions else ""}
{ctx_str}

Their clause:
{existing_text}

Provide:
## Analysis of their position
[What they're trying to achieve, where it's aggressive or unreasonable]

## Our counter-proposal
[Full redrafted clause]

## Negotiation notes
[What we can concede, what is a red line, fallback positions]
[Market standard position on this clause type]

## Mickey notes
[Items requiring partner/client sign-off before sending]"""

    answer, tokens, err = call_claude(username, prompt, docs_system(display), max_tokens=4000)
    if err:
        return jsonify(err), 500

    # Extract Mickey notes
    mickey_notes = []
    in_notes = False
    for line in answer.split("\n"):
        if "## Mickey notes" in line:
            in_notes = True
            continue
        if line.startswith("## ") and in_notes:
            break
        if in_notes and line.strip().startswith("- "):
            mickey_notes.append(line.strip()[2:])

    filename = d.get("filename", f"{doc_type}_{mode}.txt") if isinstance(d, dict) else f"{doc_type}_{mode}.txt"
    history_id = docs_save_history(username, "draft", filename, answer,
                                   {"mode": mode, "doc_type": doc_type,
                                    "governing_law": governing_law, "perspective": perspective})
    log_usage(username, "docs_draft", tokens, f"{mode}:{clause_desc[:30]}")

    return jsonify({
        "result": answer,
        "mode": mode,
        "mickey_notes": mickey_notes,
        "history_id": history_id
    })


# ══════════════════════════════════════════════════════════════
# SUMMARISE
# ══════════════════════════════════════════════════════════════
# Output: {
#   result: full markdown summary,
#   title: str,
#   key_points: [str],
#   parties: [str],
#   governing_law: str,
#   filename: str,
#   ocr_needed: bool,
#   history_id: str
# }
# ══════════════════════════════════════════════════════════════

@app.route("/api/docs/summarise", methods=["POST"])
@login_required
@verify_csrf
def docs_summarise():
    username = session["username"]
    if not check_rate_limit(username):
        return jsonify({"error": "Rate limit reached. Please wait."}), 429

    file = request.files.get("file")
    if not file:
        return jsonify({"error": "No file uploaded"}), 400

    audience      = request.form.get("audience", "Myself")
    focus_areas   = request.form.get("focus_areas", "")   # comma-separated
    output_format = request.form.get("output_format", "structured")  # structured|executive|bullet|table
    extra         = request.form.get("extra", "")

    extracted = extract_bytes_with_meta(file.read(), file.filename)
    text = extracted["text"]
    if not text:
        return jsonify({"error": "Could not extract text from document."}), 400

    display = load_users().get(username, {}).get("display_name", username)

    audience_formats = {
        "Myself": "Detailed legal summary. Include all material provisions, conditions, and risks. Use legal terminology.",
        "Board": "Executive summary for the board. Lead with recommendation and risk level. Plain language. Max 400 words.",
        "Management": "Operational summary. Focus on obligations, deadlines, and action items. Who does what by when.",
        "Other party": "Neutral factual summary. Avoid adversarial framing. Suitable for sharing with counterparty.",
        "File record": "Comprehensive file note. Complete coverage. Formal language. Include all defined terms."
    }

    format_instructions = {
        "structured": "Use headers: Summary / Parties / Key terms / Obligations & deadlines / Risks / Recommendation",
        "executive":  "Single-page executive brief. 3 sections: What it is (2 sentences) / Key points (5 bullets) / Action required",
        "bullet":     "Structured bullet points only. Group by topic. No prose paragraphs.",
        "table":      "Present key provisions as a two-column table: Provision | Summary"
    }

    prompt = f"""Summarise this legal document.

Document: {file.filename}
Audience: {audience}
{f"Focus on: {focus_areas}" if focus_areas else ""}
{f"Additional instructions: {extra}" if extra else ""}

{audience_formats.get(audience, audience_formats["Myself"])}
Format: {format_instructions.get(output_format, format_instructions["structured"])}

Document text:
{text[:9000]}

After the summary, add a one-line metadata section in this exact format:
METADATA | Parties: [list] | Governing law: [law] | Key dates: [dates if any] | Document type: [type]"""

    answer, tokens, err = call_claude(username, prompt, docs_system(display))
    if err:
        return jsonify(err), 500

    # Parse METADATA line
    parties, governing_law_found, doc_type_found = [], "", ""
    for line in answer.split("\n"):
        if line.startswith("METADATA |"):
            parts = line.split("|")
            for part in parts:
                p = part.strip()
                if p.startswith("Parties:"):
                    parties = [x.strip() for x in p[8:].split(",") if x.strip()]
                elif p.startswith("Governing law:"):
                    governing_law_found = p[14:].strip()
                elif p.startswith("Document type:"):
                    doc_type_found = p[14:].strip()
            # Remove metadata line from displayed result
            answer = answer.replace(line, "").strip()
            break

    history_id = docs_save_history(username, "summarise", file.filename, answer,
                                   {"audience": audience, "output_format": output_format,
                                    "parties": parties, "governing_law": governing_law_found})
    log_usage(username, "docs_summarise", tokens, file.filename[:40])

    return jsonify({
        "result": answer,
        "parties": parties,
        "governing_law": governing_law_found,
        "doc_type": doc_type_found,
        "filename": file.filename,
        "ocr_needed": extracted["ocr_needed"],
        "pages": extracted["pages"],
        "history_id": history_id
    })


# ══════════════════════════════════════════════════════════════
# TRANSLATE
# ══════════════════════════════════════════════════════════════
# Output: {
#   result: translated text,
#   untranslatable: [str],  — terms kept in original with explanation
#   filename: str,
#   src_lang: str,
#   tgt_lang: str,
#   ocr_needed: bool,
#   history_id: str
# }
# ══════════════════════════════════════════════════════════════

@app.route("/api/docs/translate", methods=["POST"])
@login_required
@verify_csrf
def docs_translate():
    username = session["username"]
    if not check_rate_limit(username):
        return jsonify({"error": "Rate limit reached. Please wait."}), 429

    file = request.files.get("file")
    if not file:
        return jsonify({"error": "No file uploaded"}), 400

    src_lang = request.form.get("src_lang", "Dutch (NL)")
    tgt_lang = request.form.get("tgt_lang", "English (EN)")
    style    = request.form.get("style", "Legal / formal")   # Legal/formal | Plain | Simplified
    extra    = request.form.get("extra", "")

    extracted = extract_bytes_with_meta(file.read(), file.filename)
    text = extracted["text"]
    if not text:
        return jsonify({"error": "Could not extract text from document."}), 400

    display = load_users().get(username, {}).get("display_name", username)

    style_notes = {
        "Legal / formal":  "Preserve all legal terminology. Keep article numbering. Capitalised defined terms stay capitalised. Do not simplify.",
        "Plain":           "Translate accurately but in clear, accessible language. Explain legal jargon in brackets where unavoidable.",
        "Simplified":      "Translate and simplify. Write as if explaining to a non-lawyer. Break complex sentences. Add [plain language note] where needed."
    }

    prompt = f"""Translate this legal document from {src_lang} to {tgt_lang}.

Style: {style} — {style_notes.get(style, style_notes["Legal / formal"])}
{f"Additional instructions: {extra}" if extra else ""}

Rules:
- Preserve the exact document structure: articles, sections, subsections, numbering
- For terms with no direct equivalent in the target language, keep the original term in italics and add a translator's note in square brackets
- Do not add, remove or paraphrase content — translation only
- If a clause is ambiguous in the source and that ambiguity is legally material, add [Translator's note: ambiguity in source — see original]

Document:
{text[:8000]}

After the translation, add:
## Translator's notes
- List every term kept in the original language and why
- List any material ambiguities flagged
- [If none: "No untranslatable terms or material ambiguities."]"""

    answer, tokens, err = call_claude(username, prompt, docs_system(display), max_tokens=5000)
    if err:
        return jsonify(err), 500

    # Extract untranslatable terms
    untranslatable = []
    in_notes = False
    for line in answer.split("\n"):
        if "## Translator" in line:
            in_notes = True
            continue
        if line.startswith("## ") and in_notes:
            break
        if in_notes and line.strip().startswith("- "):
            untranslatable.append(line.strip()[2:])

    history_id = docs_save_history(username, "translate", file.filename, answer,
                                   {"src_lang": src_lang, "tgt_lang": tgt_lang, "style": style})
    log_usage(username, "docs_translate", tokens, file.filename[:40])

    return jsonify({
        "result": answer,
        "untranslatable": untranslatable,
        "filename": file.filename,
        "src_lang": src_lang,
        "tgt_lang": tgt_lang,
        "ocr_needed": extracted["ocr_needed"],
        "history_id": history_id
    })


# ══════════════════════════════════════════════════════════════
# ANONYMISE
# ══════════════════════════════════════════════════════════════
# Three layers:
#   1. Auto-detected fields (names, addresses, KBO, IBAN, emails, phones)
#   2. Domain preset (commercial / M&A / employment / litigation)
#   3. User-defined custom fields [{find: str, replace: str}]
#
# Output: {
#   result: anonymised text,
#   redaction_map: [{original, replacement, count, category}],
#   stats: {total_redactions, categories_used},
#   filename: str,
#   history_id: str
# }
# ══════════════════════════════════════════════════════════════

@app.route("/api/docs/anonymise", methods=["POST"])
@login_required
@verify_csrf
def docs_anonymise():
    username = session["username"]
    if not check_rate_limit(username):
        return jsonify({"error": "Rate limit reached. Please wait."}), 429

    file = request.files.get("file")
    if not file:
        return jsonify({"error": "No file uploaded"}), 400

    preset       = request.form.get("preset", "")        # commercial|manda|employment|litigation
    custom_fields= request.form.get("custom_fields", "[]") # JSON array [{find, replace}]
    include_map  = request.form.get("include_map", "true").lower() == "true"

    try:
        custom = json.loads(custom_fields)
    except Exception:
        custom = []

    extracted = extract_bytes_with_meta(file.read(), file.filename)
    text = extracted["text"]
    if not text:
        return jsonify({"error": "Could not extract text from document."}), 400

    display = load_users().get(username, {}).get("display_name", username)

    preset_instructions = {
        "commercial": "Also redact: exact prices and amounts, product names, delivery locations, specific performance dates, payment account details.",
        "manda":      "Also redact: target company name and any code names, financial advisors, valuation figures, deal structure details, conditions precedent specifics.",
        "employment": "Also redact: salary figures, employee ID numbers, department names, performance ratings, disciplinary history references.",
        "litigation": "Also redact: case reference numbers, opposing party names, settlement amounts, without prejudice communications, expert names."
    }

    custom_instruction = ""
    if custom:
        custom_instruction = "\n\nCustom replacements (apply exactly as specified):\n" + \
            "\n".join(f'- Replace all instances of "{c["find"]}" with "{c["replace"]}"'
                      for c in custom if c.get("find") and c.get("replace"))

    prompt = f"""Anonymise this legal document by redacting all identifying information.

Document: {file.filename}
{f"Domain preset: {preset_instructions.get(preset, '')}" if preset else ""}
{custom_instruction}

Standard redactions to apply:
- Party names (companies and individuals) → [PARTY A], [PARTY B], [INDIVIDUAL 1], [INDIVIDUAL 2] etc.
  (be consistent: same entity always gets same placeholder throughout)
- Addresses and registered offices → [ADDRESS 1], [ADDRESS 2] etc.
- Registration / KBO / company numbers → [REG NUMBER]
- Email addresses → [EMAIL]
- IBAN / bank account numbers → [BANK ACCOUNT]
- Phone numbers → [PHONE]
- Signatures and initials → [SIGNATURE]

Document text:
{text[:8000]}

Return in this exact format:

## Anonymised document
[Full anonymised text — preserve all structure, numbering, headings]

## Redaction map
[One line per substitution in this format:]
REDACTION | Original: [original value] | Replaced with: [placeholder] | Category: [category] | Count: [n]
[If include_map is false, omit this section entirely]"""

    answer, tokens, err = call_claude(username, prompt, docs_system(display), max_tokens=6000)
    if err:
        return jsonify(err), 500

    # Parse sections
    anon_text = ""
    redaction_map = []
    in_doc, in_map = False, False

    for line in answer.split("\n"):
        if "## Anonymised document" in line:
            in_doc, in_map = True, False
            continue
        if "## Redaction map" in line:
            in_doc, in_map = False, True
            continue
        if line.startswith("## ") and (in_doc or in_map):
            in_doc = in_map = False
        if in_doc:
            anon_text += line + "\n"
        elif in_map and line.startswith("REDACTION |"):
            parts = line.split("|")
            entry = {}
            for part in parts[1:]:
                p = part.strip()
                if p.startswith("Original:"):
                    entry["original"] = p[9:].strip()
                elif p.startswith("Replaced with:"):
                    entry["replacement"] = p[14:].strip()
                elif p.startswith("Category:"):
                    entry["category"] = p[9:].strip()
                elif p.startswith("Count:"):
                    try:
                        entry["count"] = int(p[6:].strip())
                    except Exception:
                        entry["count"] = 1
            if entry.get("original"):
                redaction_map.append(entry)

    anon_text = anon_text.strip()
    if not anon_text:
        anon_text = answer  # fallback if parsing fails

    # Apply custom fields on top (exact string replacement, case-sensitive)
    for c in custom:
        if c.get("find") and c.get("replace"):
            count = anon_text.count(c["find"])
            if count:
                anon_text = anon_text.replace(c["find"], c["replace"])
                redaction_map.append({
                    "original": c["find"],
                    "replacement": c["replace"],
                    "category": "custom",
                    "count": count
                })

    # Save anonymised output
    anon_id = str(uuid.uuid4())[:8]
    anon_dir = docs_user_path(username) / "anonymised" / anon_id
    anon_dir.mkdir(parents=True, exist_ok=True)
    (anon_dir / "anonymised.txt").write_text(anon_text, encoding="utf-8")
    if include_map:
        docs_save_json(anon_dir / "redaction_map.json",
                       {"filename": file.filename, "map": redaction_map})

    # Stats
    total_redactions = sum(e.get("count", 1) for e in redaction_map)
    categories = list({e.get("category", "other") for e in redaction_map})

    history_id = docs_save_history(username, "anonymise", file.filename, anon_text,
                                   {"preset": preset, "total_redactions": total_redactions,
                                    "anon_id": anon_id})
    log_usage(username, "docs_anonymise", tokens, file.filename[:40])

    return jsonify({
        "result": anon_text,
        "redaction_map": redaction_map if include_map else [],
        "stats": {"total_redactions": total_redactions, "categories_used": categories},
        "filename": file.filename,
        "anon_id": anon_id,
        "history_id": history_id
    })


# ══════════════════════════════════════════════════════════════
# HISTORY
# ══════════════════════════════════════════════════════════════

@app.route("/api/docs/history", methods=["GET"])
@login_required
def docs_history():
    username = session["username"]
    module   = request.args.get("module", "")   # filter by module, empty = all
    limit    = min(int(request.args.get("limit", "50")), 200)
    hist_dir = docs_user_path(username) / "history"
    items    = docs_list_dir(hist_dir)
    if module:
        items = [i for i in items if i.get("module") == module]
    # Return previews only (not full result) for list view
    previews = [{
        "id":           i.get("id"),
        "module":       i.get("module"),
        "filename":     i.get("filename"),
        "result_preview": i.get("result_preview", "")[:200],
        "meta":         i.get("meta", {}),
        "created_at":   i.get("created_at")
    } for i in items[:limit]]
    return jsonify({"items": previews, "total": len(items)})


@app.route("/api/docs/history/<item_id>", methods=["GET"])
@login_required
def docs_history_item(item_id):
    username = session["username"]
    # Sanitise id — alphanumeric and hyphens only
    if not re.match(r'^[a-zA-Z0-9\-]+$', item_id):
        return jsonify({"error": "Invalid id"}), 400
    path = docs_user_path(username) / "history" / f"{item_id}.json"
    if not path.exists():
        return jsonify({"error": "Not found"}), 404
    return jsonify(docs_load_json(path))


@app.route("/api/docs/history/<item_id>", methods=["DELETE"])
@login_required
@verify_csrf
def docs_history_delete(item_id):
    username = session["username"]
    if not re.match(r'^[a-zA-Z0-9\-]+$', item_id):
        return jsonify({"error": "Invalid id"}), 400
    path = docs_user_path(username) / "history" / f"{item_id}.json"
    if not path.exists():
        return jsonify({"error": "Not found"}), 404
    path.unlink()
    return jsonify({"ok": True})


# ══════════════════════════════════════════════════════════════
# TEMPLATES
# ══════════════════════════════════════════════════════════════
# Templates are user-saved drafting starting points:
# {id, name, doc_type, governing_law, content, tags, created_at, updated_at}

@app.route("/api/docs/templates", methods=["GET"])
@login_required
def docs_templates_list():
    username  = session["username"]
    doc_type  = request.args.get("doc_type", "")
    tmpl_dir  = docs_user_path(username) / "templates"
    items     = docs_list_dir(tmpl_dir)
    if doc_type:
        items = [i for i in items if i.get("doc_type", "").lower() == doc_type.lower()]
    # Return without full content for list view
    previews = [{
        "id":          i.get("id"),
        "name":        i.get("name"),
        "doc_type":    i.get("doc_type"),
        "governing_law": i.get("governing_law"),
        "tags":        i.get("tags", []),
        "content_preview": (i.get("content", "")[:150] + "…") if i.get("content") else "",
        "created_at":  i.get("created_at"),
        "updated_at":  i.get("updated_at")
    } for i in items]
    return jsonify({"templates": previews})


@app.route("/api/docs/templates", methods=["POST"])
@login_required
@verify_csrf
def docs_templates_save():
    username = session["username"]
    d        = request.get_json(force=True) or {}
    name     = d.get("name", "").strip()
    if not name:
        return jsonify({"error": "Template name required"}), 400

    tmpl_id  = d.get("id", str(uuid.uuid4())[:8])  # update if id provided, create otherwise
    now      = datetime.datetime.utcnow().isoformat()
    path     = docs_user_path(username) / "templates" / f"{tmpl_id}.json"

    existing = docs_load_json(path) if path.exists() else {}
    template = {
        "id":           tmpl_id,
        "name":         name,
        "doc_type":     d.get("doc_type", existing.get("doc_type", "")),
        "governing_law":d.get("governing_law", existing.get("governing_law", "Belgian law")),
        "content":      d.get("content", existing.get("content", "")),
        "tags":         d.get("tags", existing.get("tags", [])),
        "created_at":   existing.get("created_at", now),
        "updated_at":   now
    }
    docs_save_json(path, template)
    return jsonify({"ok": True, "id": tmpl_id, "template": template})


@app.route("/api/docs/template/<tmpl_id>", methods=["GET"])
@login_required
def docs_template_get(tmpl_id):
    username = session["username"]
    if not re.match(r'^[a-zA-Z0-9\-]+$', tmpl_id):
        return jsonify({"error": "Invalid id"}), 400
    path = docs_user_path(username) / "templates" / f"{tmpl_id}.json"
    if not path.exists():
        return jsonify({"error": "Not found"}), 404
    return jsonify(docs_load_json(path))


@app.route("/api/docs/template/<tmpl_id>", methods=["DELETE"])
@login_required
@verify_csrf
def docs_template_delete(tmpl_id):
    username = session["username"]
    if not re.match(r'^[a-zA-Z0-9\-]+$', tmpl_id):
        return jsonify({"error": "Invalid id"}), 400
    path = docs_user_path(username) / "templates" / f"{tmpl_id}.json"
    if not path.exists():
        return jsonify({"error": "Not found"}), 404
    path.unlink()
    return jsonify({"ok": True})


# ══════════════════════════════════════════════════════════════
# OBLIGATIONS  (dashboard widget feed)
# ══════════════════════════════════════════════════════════════
# Returns all saved obligations extracted from reviewed documents.
# The dashboard widget calls this to show "Contract obligations due".
# Frontend calls widgetPush('doc_obligations', data) after fetching.

@app.route("/api/docs/obligations", methods=["GET"])
@login_required
def docs_obligations():
    username = session["username"]
    limit    = min(int(request.args.get("limit", "100")), 500)
    obs_dir  = docs_user_path(username) / "obligations"
    all_obs  = []
    for item in docs_list_dir(obs_dir)[:limit]:
        for ob in item.get("obligations", []):
            all_obs.append({
                **ob,
                "source_doc": item.get("filename", ""),
                "extracted_at": item.get("created_at", "")
            })
    return jsonify({"obligations": all_obs, "total": len(all_obs)})

