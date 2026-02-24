"""
compliance/routes.py
Mickey Compliance Module — Flask Blueprint
All routes follow /api/compliance/{register}/{action} convention.
"""
import os
import json
from functools import wraps
from io import BytesIO
from flask import Blueprint, request, jsonify, session, render_template, redirect, url_for, send_file
from datetime import datetime
from werkzeug.utils import secure_filename

from . import data as cd

compliance_bp = Blueprint("compliance", __name__, template_folder="../templates/compliance")

UPLOAD_ROOT = "/opt/mickey/data/uploads"
ALLOWED_EXTENSIONS = {"pdf", "docx", "doc", "xlsx", "txt", "md"}
MAX_FILE_MB = 20

# ── Auth decorators ───────────────────────────────────────────────────────────

def login_required(f):
    @wraps(f)
    def decorated(*args, **kwargs):
        if "firm_id" not in session or "user_id" not in session:
            return jsonify({"error": "Authentication required"}), 401
        return f(*args, **kwargs)
    return decorated

def module_access_required(module_name):
    """Check user has access to this compliance module."""
    def decorator(f):
        @wraps(f)
        def decorated(*args, **kwargs):
            if "firm_id" not in session:
                return jsonify({"error": "Authentication required"}), 401
            modules = session.get("modules", [])
            if module_name not in modules:
                return jsonify({"error": "Access denied to this module"}), 403
            return f(*args, **kwargs)
        return decorated
    return decorator

def editor_required(f):
    """Require editor or admin role."""
    @wraps(f)
    def decorated(*args, **kwargs):
        role = session.get("compliance_role", "viewer")
        if role not in ("admin", "editor"):
            return jsonify({"error": "Editor access required"}), 403
        return f(*args, **kwargs)
    return decorated

def admin_required(f):
    """Require admin role."""
    @wraps(f)
    def decorated(*args, **kwargs):
        role = session.get("compliance_role", "viewer")
        if role != "admin":
            return jsonify({"error": "Admin access required"}), 403
        return f(*args, **kwargs)
    return decorated

def get_firm():
    return session.get("firm_id")

def ok(data=None):
    if data is None:
        return jsonify({"ok": True})
    return jsonify({"ok": True, "data": data})

def err(msg, code=400):
    return jsonify({"error": msg}), code

# ── Main compliance page ──────────────────────────────────────────────────────

@compliance_bp.route("/compliance")
@login_required
def index():
    return redirect("/app")

# ── Compliance settings ───────────────────────────────────────────────────────

@compliance_bp.route("/api/compliance/settings", methods=["GET"])
@login_required
@module_access_required("compliance")
def get_settings():
    return ok(cd.get_compliance_settings(get_firm()))

@compliance_bp.route("/api/compliance/settings", methods=["POST"])
@login_required
@module_access_required("compliance")
@admin_required
def save_settings():
    data = request.get_json(silent=True) or {}
    return ok(cd.save_compliance_settings(get_firm(), data))

# ── DPO ──────────────────────────────────────────────────────────────────────

@compliance_bp.route("/api/compliance/dpo", methods=["GET"])
@login_required
@module_access_required("compliance")
def get_dpo():
    return ok(cd.get_dpo(get_firm()))

@compliance_bp.route("/api/compliance/dpo", methods=["POST"])
@login_required
@module_access_required("compliance")
@editor_required
def save_dpo():
    data = request.get_json(silent=True) or {}
    return ok(cd.save_dpo(get_firm(), data))

# ── RoPA ─────────────────────────────────────────────────────────────────────

@compliance_bp.route("/api/compliance/ropa", methods=["GET"])
@login_required
@module_access_required("compliance")
def get_ropa():
    return ok(cd.get_ropa(get_firm()))

@compliance_bp.route("/api/compliance/ropa", methods=["POST"])
@login_required
@module_access_required("compliance")
@editor_required
def add_ropa():
    data = request.get_json(silent=True) or {}
    required = ["activity_name", "role"]
    missing = [f for f in required if not data.get(f)]
    if missing:
        return err(f"Missing required fields: {', '.join(missing)}")
    return ok(cd.add_ropa(get_firm(), data)), 201

@compliance_bp.route("/api/compliance/ropa/<record_id>", methods=["PUT"])
@login_required
@module_access_required("compliance")
@editor_required
def update_ropa(record_id):
    data = request.get_json(silent=True) or {}
    result = cd.update_ropa(get_firm(), record_id, data)
    if not result:
        return err("Record not found", 404)
    return ok(result)

@compliance_bp.route("/api/compliance/ropa/<record_id>", methods=["DELETE"])
@login_required
@module_access_required("compliance")
@editor_required
def delete_ropa(record_id):
    cd.delete_ropa(get_firm(), record_id)
    return ok()

# ── DSR ──────────────────────────────────────────────────────────────────────

@compliance_bp.route("/api/compliance/dsrs", methods=["GET"])
@login_required
@module_access_required("compliance")
def get_dsrs():
    return ok(cd.get_dsrs(get_firm()))

@compliance_bp.route("/api/compliance/dsrs", methods=["POST"])
@login_required
@module_access_required("compliance")
@editor_required
def add_dsr():
    data = request.get_json(silent=True) or {}
    required = ["requester_name", "right", "received_date", "role"]
    missing = [f for f in required if not data.get(f)]
    if missing:
        return err(f"Missing required fields: {', '.join(missing)}")
    return ok(cd.add_dsr(get_firm(), data)), 201

@compliance_bp.route("/api/compliance/dsrs/<record_id>", methods=["PUT"])
@login_required
@module_access_required("compliance")
@editor_required
def update_dsr(record_id):
    data = request.get_json(silent=True) or {}
    result = cd.update_dsr(get_firm(), record_id, data)
    if not result:
        return err("Record not found", 404)
    return ok(result)

@compliance_bp.route("/api/compliance/dsrs/<record_id>/close", methods=["POST"])
@login_required
@module_access_required("compliance")
@editor_required
def close_dsr(record_id):
    data = request.get_json(silent=True) or {}
    result = cd.close_dsr(get_firm(), record_id, data.get("outcome", ""))
    if not result:
        return err("Record not found", 404)
    return ok(result)

# ── Data Breach ───────────────────────────────────────────────────────────────

@compliance_bp.route("/api/compliance/breaches", methods=["GET"])
@login_required
@module_access_required("compliance")
def get_breaches():
    return ok(cd.get_breaches(get_firm()))

@compliance_bp.route("/api/compliance/breaches", methods=["POST"])
@login_required
@module_access_required("compliance")
@editor_required
def add_breach():
    data = request.get_json(silent=True) or {}
    required = ["description", "detected_at", "breach_type"]
    missing = [f for f in required if not data.get(f)]
    if missing:
        return err(f"Missing required fields: {', '.join(missing)}")
    return ok(cd.add_breach(get_firm(), data)), 201

@compliance_bp.route("/api/compliance/breaches/<record_id>", methods=["PUT"])
@login_required
@module_access_required("compliance")
@editor_required
def update_breach(record_id):
    data = request.get_json(silent=True) or {}
    result = cd.update_breach(get_firm(), record_id, data)
    if not result:
        return err("Record not found", 404)
    return ok(result)

# ── DPIA ─────────────────────────────────────────────────────────────────────

@compliance_bp.route("/api/compliance/dpias", methods=["GET"])
@login_required
@module_access_required("compliance")
def get_dpias():
    return ok(cd.get_dpias(get_firm()))

@compliance_bp.route("/api/compliance/dpias", methods=["POST"])
@login_required
@module_access_required("compliance")
@editor_required
def add_dpia():
    data = request.get_json(silent=True) or {}
    if not data.get("activity_name"):
        return err("Activity name required")
    return ok(cd.add_dpia(get_firm(), data)), 201

@compliance_bp.route("/api/compliance/dpias/<record_id>/advance", methods=["POST"])
@login_required
@module_access_required("compliance")
@editor_required
def advance_dpia(record_id):
    result = cd.advance_dpia(get_firm(), record_id)
    if not result:
        return err("Record not found", 404)
    return ok(result)

@compliance_bp.route("/api/compliance/dpias/<record_id>", methods=["PUT"])
@login_required
@module_access_required("compliance")
@editor_required
def update_dpia(record_id):
    data = request.get_json(silent=True) or {}
    result = cd.update_dpia(get_firm(), record_id, data)
    if not result:
        return err("Record not found", 404)
    return ok(result)

# ── TIA ──────────────────────────────────────────────────────────────────────

@compliance_bp.route("/api/compliance/tias", methods=["GET"])
@login_required
@module_access_required("compliance")
def get_tias():
    return ok(cd.get_tias(get_firm()))

@compliance_bp.route("/api/compliance/tias", methods=["POST"])
@login_required
@module_access_required("compliance")
@editor_required
def add_tia():
    data = request.get_json(silent=True) or {}
    required = ["importer_name", "destination_country", "transfer_mechanism"]
    missing = [f for f in required if not data.get(f)]
    if missing:
        return err(f"Missing required fields: {', '.join(missing)}")
    return ok(cd.add_tia(get_firm(), data)), 201

@compliance_bp.route("/api/compliance/tias/<record_id>", methods=["PUT"])
@login_required
@module_access_required("compliance")
@editor_required
def update_tia(record_id):
    data = request.get_json(silent=True) or {}
    result = cd.update_tia(get_firm(), record_id, data)
    if not result:
        return err("Record not found", 404)
    return ok(result)

@compliance_bp.route("/api/compliance/tias/<record_id>", methods=["DELETE"])
@login_required
@module_access_required("compliance")
@editor_required
def delete_tia(record_id):
    cd.delete_tia(get_firm(), record_id)
    return ok()

# ── DPA ──────────────────────────────────────────────────────────────────────

@compliance_bp.route("/api/compliance/dpas", methods=["GET"])
@login_required
@module_access_required("compliance")
def get_dpas():
    return ok(cd.get_dpas(get_firm()))

@compliance_bp.route("/api/compliance/dpas", methods=["POST"])
@login_required
@module_access_required("compliance")
@editor_required
def add_dpa():
    data = request.get_json(silent=True) or {}
    required = ["counterparty", "our_role"]
    missing = [f for f in required if not data.get(f)]
    if missing:
        return err(f"Missing required fields: {', '.join(missing)}")
    return ok(cd.add_dpa(get_firm(), data)), 201

@compliance_bp.route("/api/compliance/dpas/<record_id>", methods=["PUT"])
@login_required
@module_access_required("compliance")
@editor_required
def update_dpa(record_id):
    data = request.get_json(silent=True) or {}
    result = cd.update_dpa(get_firm(), record_id, data)
    if not result:
        return err("Record not found", 404)
    return ok(result)

@compliance_bp.route("/api/compliance/dpas/<record_id>", methods=["DELETE"])
@login_required
@module_access_required("compliance")
@editor_required
def delete_dpa(record_id):
    cd.delete_dpa(get_firm(), record_id)
    return ok()

# ── Gifts ─────────────────────────────────────────────────────────────────────

@compliance_bp.route("/api/compliance/gifts", methods=["GET"])
@login_required
@module_access_required("compliance")
def get_gifts():
    return ok(cd.get_gifts(get_firm()))

@compliance_bp.route("/api/compliance/gifts", methods=["POST"])
@login_required
@module_access_required("compliance")
@editor_required
def add_gift():
    data = request.get_json(silent=True) or {}
    required = ["employee_name", "direction", "gift_type", "third_party", "value", "gift_date"]
    missing = [f for f in required if not data.get(f)]
    if missing:
        return err(f"Missing required fields: {', '.join(missing)}")
    return ok(cd.add_gift(get_firm(), data)), 201

@compliance_bp.route("/api/compliance/gifts/<record_id>/approve", methods=["POST"])
@login_required
@module_access_required("compliance")
@editor_required
def approve_gift(record_id):
    data = request.get_json(silent=True) or {}
    result = cd.update_gift_status(get_firm(), record_id, "approved", data.get("note", ""))
    if not result:
        return err("Record not found", 404)
    return ok(result)

@compliance_bp.route("/api/compliance/gifts/<record_id>/decline", methods=["POST"])
@login_required
@module_access_required("compliance")
@editor_required
def decline_gift(record_id):
    data = request.get_json(silent=True) or {}
    result = cd.update_gift_status(get_firm(), record_id, "declined", data.get("note", ""))
    if not result:
        return err("Record not found", 404)
    return ok(result)

# ── COI ──────────────────────────────────────────────────────────────────────

@compliance_bp.route("/api/compliance/coi", methods=["GET"])
@login_required
@module_access_required("compliance")
def get_coi():
    return ok(cd.get_coi(get_firm()))

@compliance_bp.route("/api/compliance/coi", methods=["POST"])
@login_required
@module_access_required("compliance")
@editor_required
def add_coi():
    data = request.get_json(silent=True) or {}
    required = ["person_name", "person_role", "declaration_year"]
    missing = [f for f in required if not data.get(f)]
    if missing:
        return err(f"Missing required fields: {', '.join(missing)}")
    return ok(cd.add_coi(get_firm(), data)), 201

@compliance_bp.route("/api/compliance/coi/<record_id>", methods=["PUT"])
@login_required
@module_access_required("compliance")
@editor_required
def update_coi(record_id):
    data = request.get_json(silent=True) or {}
    result = cd.update_coi(get_firm(), record_id, data)
    if not result:
        return err("Record not found", 404)
    return ok(result)

# ── Third-Party DD ────────────────────────────────────────────────────────────

@compliance_bp.route("/api/compliance/tpdd", methods=["GET"])
@login_required
@module_access_required("compliance")
def get_tpdd():
    return ok(cd.get_tpdd(get_firm()))

@compliance_bp.route("/api/compliance/tpdd", methods=["POST"])
@login_required
@module_access_required("compliance")
@editor_required
def add_tpdd():
    data = request.get_json(silent=True) or {}
    required = ["company_name", "third_party_type", "country"]
    missing = [f for f in required if not data.get(f)]
    if missing:
        return err(f"Missing required fields: {', '.join(missing)}")
    return ok(cd.add_tpdd(get_firm(), data)), 201

@compliance_bp.route("/api/compliance/tpdd/<record_id>", methods=["PUT"])
@login_required
@module_access_required("compliance")
@editor_required
def update_tpdd(record_id):
    data = request.get_json(silent=True) or {}
    result = cd.update_tpdd(get_firm(), record_id, data)
    if not result:
        return err("Record not found", 404)
    return ok(result)

# ── Donations ─────────────────────────────────────────────────────────────────

@compliance_bp.route("/api/compliance/donations", methods=["GET"])
@login_required
@module_access_required("compliance")
def get_donations():
    return ok(cd.get_donations(get_firm()))

@compliance_bp.route("/api/compliance/donations", methods=["POST"])
@login_required
@module_access_required("compliance")
@editor_required
def add_donation():
    data = request.get_json(silent=True) or {}
    required = ["recipient", "donation_type", "amount", "donation_date"]
    missing = [f for f in required if not data.get(f)]
    if missing:
        return err(f"Missing required fields: {', '.join(missing)}")
    return ok(cd.add_donation(get_firm(), data)), 201

@compliance_bp.route("/api/compliance/donations/<record_id>", methods=["PUT"])
@login_required
@module_access_required("compliance")
@editor_required
def update_donation(record_id):
    data = request.get_json(silent=True) or {}
    result = cd.update_donation(get_firm(), record_id, data)
    if not result:
        return err("Record not found", 404)
    return ok(result)

# ── Red Flags ─────────────────────────────────────────────────────────────────

@compliance_bp.route("/api/compliance/flags", methods=["GET"])
@login_required
@module_access_required("compliance")
def get_flags():
    # Admin + editor only regardless of module setting
    role = session.get("compliance_role", "viewer")
    if role not in ("admin", "editor"):
        return err("Access restricted", 403)
    return ok(cd.get_flags(get_firm()))

@compliance_bp.route("/api/compliance/flags", methods=["POST"])
@login_required
@module_access_required("compliance")
@editor_required
def add_flag():
    data = request.get_json(silent=True) or {}
    required = ["description", "identified_date", "severity"]
    missing = [f for f in required if not data.get(f)]
    if missing:
        return err(f"Missing required fields: {', '.join(missing)}")
    return ok(cd.add_flag(get_firm(), data)), 201

@compliance_bp.route("/api/compliance/flags/<record_id>", methods=["PUT"])
@login_required
@module_access_required("compliance")
@editor_required
def update_flag(record_id):
    data = request.get_json(silent=True) or {}
    result = cd.update_flag(get_firm(), record_id, data)
    if not result:
        return err("Record not found", 404)
    return ok(result)

# ── Whistleblowing ────────────────────────────────────────────────────────────

@compliance_bp.route("/api/compliance/wb", methods=["GET"])
@login_required
@module_access_required("whistleblowing")
def get_wb():
    return ok(cd.get_wb_cases(get_firm()))

@compliance_bp.route("/api/compliance/wb", methods=["POST"])
@login_required
@module_access_required("whistleblowing")
@editor_required
def add_wb():
    data = request.get_json(silent=True) or {}
    required = ["category", "received_date", "channel", "summary"]
    missing = [f for f in required if not data.get(f)]
    if missing:
        return err(f"Missing required fields: {', '.join(missing)}")
    return ok(cd.add_wb_case(get_firm(), data)), 201

@compliance_bp.route("/api/compliance/wb/<record_id>/status", methods=["POST"])
@login_required
@module_access_required("whistleblowing")
@editor_required
def update_wb_status(record_id):
    data = request.get_json(silent=True) or {}
    if not data.get("status"):
        return err("Status required")
    result = cd.update_wb_status(get_firm(), record_id, data["status"], data.get("note", ""))
    if not result:
        return err("Record not found", 404)
    return ok(result)

@compliance_bp.route("/api/compliance/wb/stats", methods=["GET"])
@login_required
@module_access_required("whistleblowing")
def wb_stats():
    return ok(cd.get_wb_annual_stats(get_firm()))

# ── Policies & Advice ─────────────────────────────────────────────────────────

@compliance_bp.route("/api/compliance/policies", methods=["GET"])
@login_required
@module_access_required("compliance")
def get_policies():
    module = request.args.get("module")
    return ok(cd.get_policies(get_firm(), module))

@compliance_bp.route("/api/compliance/policies", methods=["POST"])
@login_required
@module_access_required("compliance")
@editor_required
def add_policy():
    data = request.get_json(silent=True) or {}
    required = ["title", "doc_type", "module"]
    missing = [f for f in required if not data.get(f)]
    if missing:
        return err(f"Missing required fields: {', '.join(missing)}")
    return ok(cd.add_policy(get_firm(), data)), 201

@compliance_bp.route("/api/compliance/policies/upload", methods=["POST"])
@login_required
@module_access_required("compliance")
@editor_required
def upload_policy():
    if "file" not in request.files:
        return err("No file provided")
    f = request.files["file"]
    if not f.filename:
        return err("No file selected")
    ext = f.filename.rsplit(".", 1)[-1].lower() if "." in f.filename else ""
    if ext not in ALLOWED_EXTENSIONS:
        return err(f"File type not allowed. Use: {', '.join(ALLOWED_EXTENSIONS)}")
    firm_id = get_firm()
    upload_dir = os.path.join(UPLOAD_ROOT, firm_id, "compliance", "policies")
    os.makedirs(upload_dir, exist_ok=True)
    # Check size
    f.seek(0, 2)
    size = f.tell()
    f.seek(0)
    if size > MAX_FILE_MB * 1024 * 1024:
        return err(f"File too large. Maximum {MAX_FILE_MB}MB.")
    safe_name = secure_filename(f.filename)
    file_path = os.path.join(upload_dir, safe_name)
    f.save(file_path)
    meta = {
        "title": request.form.get("title", safe_name),
        "doc_type": request.form.get("doc_type", "policy"),
        "module": request.form.get("module", "gdpr"),
        "file_path": file_path,
        "file_name": safe_name,
        "file_ext": ext,
        "jurisdictions": request.form.get("jurisdictions", ""),
        "from_radar": request.form.get("from_radar", "false") == "true",
    }
    result = cd.add_policy(firm_id, meta)
    return ok(result), 201

@compliance_bp.route("/api/compliance/policies/<record_id>", methods=["PUT"])
@login_required
@module_access_required("compliance")
@editor_required
def update_policy(record_id):
    data = request.get_json(silent=True) or {}
    result = cd.update_policy(get_firm(), record_id, data)
    if not result:
        return err("Record not found", 404)
    return ok(result)

@compliance_bp.route("/api/compliance/policies/<record_id>", methods=["DELETE"])
@login_required
@module_access_required("compliance")
@editor_required
def delete_policy(record_id):
    cd.delete_policy(get_firm(), record_id)
    return ok()

# ── Dashboard ─────────────────────────────────────────────────────────────────

@compliance_bp.route("/api/compliance/dashboard", methods=["GET"])
@login_required
@module_access_required("compliance")
def get_dashboard():
    return ok(cd.get_dashboard(get_firm()))

# ── User role management (settings integration) ───────────────────────────────

@compliance_bp.route("/api/settings/users", methods=["GET"])
@login_required
@admin_required
def get_users():
    """Return all users in the firm with their compliance roles and modules."""
    from auth.data import load_firm
    firm = load_firm(get_firm())
    if not firm:
        return err("Firm not found", 404)
    users = []
    for uid, u in firm.get("users", {}).items():
        users.append({
            "user_id": uid,
            "display_name": u.get("display_name", ""),
            "email": u.get("email", ""),
            "status": u.get("status", "active"),
            "compliance_role": u.get("compliance_role", "viewer"),
            "modules": u.get("modules", []),
            "last_signin": u.get("last_login", ""),
            "created_at": u.get("created_at", ""),
        })
    return ok(users)

@compliance_bp.route("/api/settings/users/<user_id>", methods=["PUT"])
@login_required
@admin_required
def update_user(user_id):
    """Update a user's compliance role and module access."""
    from auth.data import load_firm, save_firm
    data = request.get_json(silent=True) or {}
    firm = load_firm(get_firm())
    if not firm or user_id not in firm.get("users", {}):
        return err("User not found", 404)
    # Prevent admin from demoting themselves
    if user_id == session.get("user_id") and data.get("compliance_role") != "admin":
        return err("You cannot change your own role")
    allowed_roles = ("admin", "editor", "viewer")
    allowed_modules = ("compliance", "whistleblowing", "documents", "corporate", "radar", "knowledge")
    if "compliance_role" in data:
        if data["compliance_role"] not in allowed_roles:
            return err("Invalid role")
        firm["users"][user_id]["compliance_role"] = data["compliance_role"]
    if "modules" in data:
        modules = [m for m in data["modules"] if m in allowed_modules]
        firm["users"][user_id]["modules"] = modules
    if "status" in data:
        if data["status"] in ("active", "inactive"):
            firm["users"][user_id]["status"] = data["status"]
    save_firm(firm)
    return ok(firm["users"][user_id])

@compliance_bp.route("/api/settings/users/invite", methods=["POST"])
@login_required
@admin_required
def invite_user():
    """Invite a new user to the firm via the standard token pipeline."""
    from auth.data import load_firm, create_invite_token
    from auth.email import send_invite_email
    data = request.get_json(silent=True) or {}
    email = (data.get("email") or "").strip().lower()
    if not email:
        return err("Email required")
    firm = load_firm(get_firm())
    if not firm:
        return err("Firm not found")
    # Check user limit
    active_users = [u for u in firm.get("users", {}).values() if u.get("status") == "active"]
    if len(active_users) >= 5:
        return err("Maximum 5 users per firm. Deactivate a user before inviting another.")
    # Check not already a member
    for u in firm.get("users", {}).values():
        if u.get("email", "").lower() == email:
            return err("This email address is already a member of your firm.")
    display_name = (data.get("display_name") or "").strip()
    role = data.get("compliance_role") or data.get("role") or "viewer"
    invited_by = session.get("display_name", "your admin")
    token = create_invite_token(
        firm_id      = firm["firm_id"],
        email        = email,
        display_name = display_name,
        role         = role,
        invited_by   = invited_by,
    )
    sent = send_invite_email(
        to_email     = email,
        display_name = display_name or email,
        firm_name    = firm["firm_name"],
        invited_by   = invited_by,
        token        = token,
    )
    if not sent:
        return err("Failed to send invite email — check Brevo config")
    return ok({"message": f"Invite sent to {email}"})

@compliance_bp.route("/api/settings/users/<user_id>/deactivate", methods=["POST"])
@login_required
@admin_required
def deactivate_user(user_id):
    from auth.data import load_firm, save_firm
    if user_id == session.get("user_id"):
        return err("You cannot deactivate your own account")
    firm = load_firm(get_firm())
    if not firm or user_id not in firm.get("users", {}):
        return err("User not found", 404)
    firm["users"][user_id]["status"] = "inactive"
    save_firm(firm)
    return ok()


# ── Export ─────────────────────────────────────────────────────────────────────

_EXPORT_REGISTERS = {
    "ropa":      ("Record of Processing Activities",        cd.get_ropa),
    "dsrs":      ("Data Subject Requests",                  cd.get_dsrs),
    "breaches":  ("Data Breaches",                          cd.get_breaches),
    "dpias":     ("Data Protection Impact Assessments",     cd.get_dpias),
    "tias":      ("Transfer Impact Assessments",            cd.get_tias),
    "dpas":      ("Data Processing Agreements",             cd.get_dpas),
    "gifts":     ("Gifts & Hospitality",                    cd.get_gifts),
    "coi":       ("Conflicts of Interest",                  cd.get_coi),
    "tpdd":      ("Third-Party Due Diligence",              cd.get_tpdd),
    "donations": ("Political Donations",                    cd.get_donations),
    "flags":     ("Red Flag Log",                           cd.get_flags),
    "wb":        ("Whistleblowing Cases",                   cd.get_wb_cases),
}

# Per-register field schema: (field_key, column_header) in display order.
_REGISTER_COLS = {
    "ropa": [
        ("activity_name","Activity Name"),("role","Role"),("purposes","Purposes"),
        ("legal_basis","Legal Basis"),("data_subjects","Data Subjects"),
        ("data_categories","Data Categories"),("recipients","Recipients"),
        ("transfers","International Transfers"),("retention","Retention Period"),
        ("special_category","Special Category"),("security_measures","Security Measures"),
        ("dpia_required","DPIA Required"),("controller_name","Controller Name"),
        ("controller_contact","Controller Contact"),("proc_categories","Processing Categories"),
        ("dpa_ref","DPA Reference"),("review_date","Review Date"),("notes","Notes"),
    ],
    "dsrs": [
        ("requester_name","Requester"),("right","Right Exercised"),("role","Our Role"),
        ("received_date","Date Received"),("deadline","Deadline"),("channel","Channel"),
        ("identity_verified","Identity Verified"),("third_party","Third Party"),
        ("status","Status"),("outcome","Outcome"),("notes","Notes"),
    ],
    "breaches": [
        ("ref","Reference"),("description","Description"),("breach_type","Type"),
        ("detected_at","Detected At"),("records_affected","Records Affected"),
        ("data_categories","Data Categories"),("risk","Risk Level"),("role","Our Role"),
        ("containment","Containment Measures"),("sa_notified_at","SA Notified At"),
        ("status","Status"),("resolution_note","Resolution Note"),("resolved_at","Resolved At"),
    ],
    "dpias": [
        ("activity_name","Activity Name"),("trigger","Trigger"),("risk","Risk Level"),
        ("ropa_link","RoPA Link"),("assessor","Assessor"),("description","Description"),
        ("stage","Current Stage"),("status","Status"),
    ],
    "tias": [
        ("exporter","Exporter"),("importer_name","Importer"),
        ("destination_country","Country"),("transfer_mechanism","Transfer Mechanism"),
        ("data_categories","Data Categories"),("purpose","Purpose"),
        ("risk","Risk"),("reassess_by","Reassess By"),
        ("supplementary_measures","Supplementary Measures"),
    ],
    "dpas": [
        ("counterparty","Counterparty"),("our_role","Our Role"),("purpose","Purpose"),
        ("signed","Date Signed"),("review_date","Review Date"),
        ("tia_ref","TIA Reference"),("notes","Notes"),
    ],
    "gifts": [
        ("employee_name","Employee"),("direction","Direction"),("gift_type","Gift Type"),
        ("third_party","Third Party"),("country","Country"),("value","Value (EUR)"),
        ("gift_date","Date"),("context","Context"),("status","Status"),("notes","Notes"),
    ],
    "coi": [
        ("person_name","Person"),("coi_category","Category"),("person_role","Role"),
        ("declaration_year","Year"),("declaration_date","Declaration Date"),
        ("conflict_status","Conflict Status"),("conflict_description","Description"),
        ("approved_by","Approved By"),("notes","Notes"),
    ],
    "tpdd": [
        ("company_name","Company"),("third_party_type","Type"),("country","Country"),
        ("dd_level","DD Level"),("last_dd","Last DD Date"),("next_review","Next Review"),
        ("status","Status"),("description","Description"),
    ],
    "donations": [
        ("recipient","Recipient"),("donation_type","Type"),("amount","Amount (EUR)"),
        ("donation_date","Date"),("jurisdiction","Jurisdiction"),
        ("board_approval_required","Board Approval Required"),
        ("justification","Justification"),("status","Status"),
    ],
    "flags": [
        ("identified_date","Date Identified"),("description","Description"),
        ("related_party","Related Party"),("country","Country"),
        ("severity","Severity"),("source","Source"),("status","Status"),
    ],
    "wb": [
        ("ref","Reference"),("received_date","Date Received"),("category","Category"),
        ("channel","Channel"),("reporter_type","Reporter Type"),("summary","Summary"),
        ("handler","Handler"),("alleged","Alleged Party"),("status","Status"),
        ("acknowledge_by","Acknowledge By"),("outcome_by","Outcome By"),
    ],
}


def _export_rows(records, register):
    """Return (headers, rows) using the defined column schema, stripping internal fields."""
    cols = _REGISTER_COLS.get(register, [])
    if not cols:
        skip = {"id", "firm_id", "created_at", "updated_at"}
        keys, seen = [], set()
        for r in records:
            for k in r:
                if k not in skip and k not in seen and not k.startswith("_"):
                    keys.append(k); seen.add(k)
        cols = [(k, k.replace("_", " ").title()) for k in keys]
    headers = [label for _, label in cols]
    rows = [[str(r.get(field) or "") for field, _ in cols] for r in records]
    return headers, rows


def _shade_cell(cell, hex_color):
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear"); shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def _build_docx(title, firm_name, headers, rows):
    from docx import Document
    from docx.shared import Pt, RGBColor, Cm
    today = datetime.now().strftime("%d %B %Y")
    doc = Document()
    for sec in doc.sections:
        sec.top_margin = Cm(2); sec.bottom_margin = Cm(2)
        sec.left_margin = Cm(2.5); sec.right_margin = Cm(2.5)
    # Title block
    h = doc.add_heading(firm_name, level=0)
    h.runs[0].font.size = Pt(22); h.runs[0].font.color.rgb = RGBColor(0x1a, 0x1a, 0x2e)
    h.paragraph_format.space_after = Pt(3)
    p = doc.add_paragraph(title)
    p.runs[0].font.size = Pt(13); p.runs[0].font.color.rgb = RGBColor(0x2d, 0x6a, 0x4f)
    p.paragraph_format.space_after = Pt(2)
    m = doc.add_paragraph(f"Exported from Mickey  ·  {today}  ·  {len(rows)} record{'s' if len(rows) != 1 else ''}")
    m.runs[0].font.size = Pt(9); m.runs[0].font.color.rgb = RGBColor(0x99, 0x99, 0x99)
    m.paragraph_format.space_after = Pt(14)
    if not rows:
        doc.add_paragraph("No records found.")
    else:
        tbl = doc.add_table(rows=1 + len(rows), cols=len(headers))
        tbl.style = "Table Grid"
        # Header row
        hrow = tbl.rows[0]
        for i, h in enumerate(headers):
            cell = hrow.cells[i]
            cell.text = h
            run = cell.paragraphs[0].runs[0] if cell.paragraphs[0].runs else cell.paragraphs[0].add_run(h)
            run.text = h; run.font.bold = True
            run.font.size = Pt(8); run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            _shade_cell(cell, "2D6A4F")
        # Data rows
        for ri, row in enumerate(rows):
            trow = tbl.rows[ri + 1]
            for ci, val in enumerate(row):
                cell = trow.cells[ci]
                cell.text = val
                run = cell.paragraphs[0].runs[0] if cell.paragraphs[0].runs else cell.paragraphs[0].add_run(val)
                run.text = val; run.font.size = Pt(8)
                if ri % 2 == 1:
                    _shade_cell(cell, "F5F5F5")
    buf = BytesIO(); doc.save(buf); buf.seek(0)
    return buf


def _build_xlsx(title, firm_name, headers, rows):
    from openpyxl import Workbook
    from openpyxl.styles import Font, PatternFill, Alignment
    from openpyxl.utils import get_column_letter
    today = datetime.now().strftime("%d %B %Y")
    wb = Workbook(); ws = wb.active
    ws.title = title[:31]
    ws["A1"] = firm_name
    ws["A1"].font = Font(name="Calibri", size=16, bold=True, color="1A1A2E")
    ws["A2"] = title
    ws["A2"].font = Font(name="Calibri", size=12, italic=True, color="2D6A4F")
    ws["A3"] = f"Exported from Mickey  ·  {today}"
    ws["A3"].font = Font(name="Calibri", size=9, color="999999")
    ws.row_dimensions[1].height = 24; ws.row_dimensions[2].height = 18; ws.row_dimensions[3].height = 14
    DATA_ROW = 5
    green_fill = PatternFill(fill_type="solid", fgColor="2D6A4F")
    alt_fill = PatternFill(fill_type="solid", fgColor="F5F5F5")
    hdr_font = Font(name="Calibri", size=10, bold=True, color="FFFFFF")
    body_font = Font(name="Calibri", size=9)
    for ci, h in enumerate(headers, 1):
        cell = ws.cell(row=DATA_ROW, column=ci, value=h)
        cell.fill = green_fill; cell.font = hdr_font
        cell.alignment = Alignment(horizontal="left", vertical="center", wrap_text=True)
    ws.row_dimensions[DATA_ROW].height = 20
    for ri, row in enumerate(rows, DATA_ROW + 1):
        fill = alt_fill if ri % 2 == 0 else None
        for ci, val in enumerate(row, 1):
            cell = ws.cell(row=ri, column=ci, value=val)
            cell.font = body_font
            cell.alignment = Alignment(vertical="top", wrap_text=True)
            if fill:
                cell.fill = fill
    for ci in range(1, len(headers) + 1):
        ws.column_dimensions[get_column_letter(ci)].width = 22
    ws.freeze_panes = ws.cell(row=DATA_ROW + 1, column=1)
    buf = BytesIO(); wb.save(buf); buf.seek(0)
    return buf


def _build_pdf(title, firm_name, headers, rows):
    from fpdf import FPDF
    today = datetime.now().strftime("%d %B %Y")
    pdf = FPDF(orientation="L", unit="mm", format="A4")
    pdf.set_margins(14, 14, 14)
    pdf.add_page()
    pdf.set_auto_page_break(auto=True, margin=14)
    GREEN, DARK, MUTED = (45, 106, 79), (26, 26, 46), (140, 140, 140)
    # Title block
    pdf.set_font("Helvetica", "B", 18); pdf.set_text_color(*DARK)
    pdf.cell(0, 9, firm_name, new_x="LMARGIN", new_y="NEXT")
    pdf.set_font("Helvetica", "B", 11); pdf.set_text_color(*GREEN)
    pdf.cell(0, 6, title, new_x="LMARGIN", new_y="NEXT")
    pdf.set_font("Helvetica", "", 8); pdf.set_text_color(*MUTED)
    pdf.cell(0, 5, f"Exported from Mickey  ·  {today}  ·  {len(rows)} record{'s' if len(rows) != 1 else ''}",
             new_x="LMARGIN", new_y="NEXT")
    pdf.ln(2)
    pdf.set_draw_color(*GREEN); pdf.set_line_width(0.5)
    pdf.line(pdf.get_x(), pdf.get_y(), pdf.w - pdf.r_margin, pdf.get_y())
    pdf.ln(4)
    if not rows:
        pdf.set_font("Helvetica", "I", 10); pdf.set_text_color(*MUTED)
        pdf.cell(0, 8, "No records found.", new_x="LMARGIN", new_y="NEXT")
    else:
        usable_w = pdf.w - pdf.l_margin - pdf.r_margin
        n = len(headers)
        col_w = max(usable_w / n, 18)
        row_h = 6
        max_ch = max(int(col_w / 1.8), 8)
        # Header
        pdf.set_fill_color(*GREEN); pdf.set_text_color(255, 255, 255)
        pdf.set_font("Helvetica", "B", 7)
        for h in headers:
            pdf.cell(col_w, row_h, h[:max_ch], border=0, fill=True)
        pdf.ln(row_h)
        # Rows
        pdf.set_font("Helvetica", "", 7)
        for i, row in enumerate(rows):
            pdf.set_fill_color(245, 245, 245) if i % 2 == 0 else pdf.set_fill_color(255, 255, 255)
            pdf.set_text_color(*DARK)
            for val in row:
                pdf.cell(col_w, row_h, str(val)[:max_ch], border=0, fill=True)
            pdf.ln(row_h)
    # Footer
    pdf.set_y(-12); pdf.set_font("Helvetica", "", 7); pdf.set_text_color(*MUTED)
    pdf.cell(0, 5, f"Mickey Legal Intelligence  ·  {firm_name}  ·  Page {pdf.page_no()}", align="C")
    return BytesIO(pdf.output())


@compliance_bp.route("/api/compliance/export/<register>", methods=["GET"])
@login_required
@module_access_required("compliance")
def export_register(register):
    if register not in _EXPORT_REGISTERS:
        return err("Unknown register", 404)
    title, getter = _EXPORT_REGISTERS[register]
    firm_id = get_firm()
    firm_name = session.get("firm_name", firm_id)
    records = getter(firm_id)
    headers, rows = _export_rows(records, register)
    fmt = request.args.get("format", "xlsx")
    today = datetime.now().strftime("%Y-%m-%d")
    base = f"mickey-{register}-{today}"
    if fmt == "xlsx":
        buf = _build_xlsx(title, firm_name, headers, rows)
        return send_file(buf, mimetype="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                         as_attachment=True, download_name=f"{base}.xlsx")
    elif fmt == "docx":
        buf = _build_docx(title, firm_name, headers, rows)
        return send_file(buf, mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                         as_attachment=True, download_name=f"{base}.docx")
    elif fmt == "pdf":
        buf = _build_pdf(title, firm_name, headers, rows)
        return send_file(buf, mimetype="application/pdf",
                         as_attachment=True, download_name=f"{base}.pdf")
    else:
        return err("Unknown format. Use docx, xlsx or pdf.", 400)
